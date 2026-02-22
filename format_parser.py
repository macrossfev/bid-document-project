#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式模板解析器
解析用户拆分后上传的投标文件格式模板 (.docx)，
识别各章节标题及其在文档中的段落位置，为模板填充提供定位依据。
"""

import re
from docx import Document
from tender_parser import classify_section, SECTION_KEYWORD_MAP


# 章节标题匹配模式（按优先级）
HEADING_PATTERNS = [
    # 一、投标函  二、报价表
    (r'^[\s]*([一二三四五六七八九十百]+)\s*[、.．]\s*(.+)', 'major'),
    # （一）投标函  （二）报价表
    (r'^[\s]*[（(]\s*([一二三四五六七八九十百]+)\s*[）)]\s*(.+)', 'sub'),
    # 第一部分 投标函
    (r'^[\s]*第\s*([一二三四五六七八九十百\d]+)\s*部分\s*(.+)', 'major'),
    # 1. 投标函  2. 报价表
    (r'^[\s]*(\d+)\s*[、.．]\s*(.+)', 'major_num'),
    # (1) 投标函
    (r'^[\s]*[（(]\s*(\d+)\s*[）)]\s*(.+)', 'sub_num'),
]

# 中文数字到阿拉伯数字映射
CN_NUM_MAP = {
    '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
    '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
    '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15,
    '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
}


def _cn_to_int(cn):
    """中文数字转阿拉伯数字"""
    if cn in CN_NUM_MAP:
        return CN_NUM_MAP[cn]
    try:
        return int(cn)
    except (ValueError, TypeError):
        return 0


def _is_heading_style(paragraph):
    """检查段落是否使用了标题样式"""
    style_name = (paragraph.style.name or '').lower()
    return 'heading' in style_name or 'title' in style_name or '标题' in style_name


def _get_paragraph_font_info(paragraph):
    """提取段落的字体格式信息（用于调试和记录）"""
    info = {
        'style': paragraph.style.name if paragraph.style else '',
        'alignment': str(paragraph.alignment) if paragraph.alignment else '',
    }
    if paragraph.runs:
        run = paragraph.runs[0]
        font = run.font
        info['font_name'] = font.name
        info['font_size'] = str(font.size) if font.size else ''
        info['bold'] = font.bold
    return info


def parse_format_template(filepath):
    """
    解析格式模板文件，识别各章节及其在文档中的位置。

    Args:
        filepath: 格式模板 .docx 文件路径

    Returns:
        dict: {
            'success': bool,
            'sections': [
                {
                    'order': int,              # 顺序号
                    'section_name': str,       # 章节名称（清理后）
                    'heading_text': str,       # 原始标题文本（用于定位）
                    'para_index': int,         # 段落在文档中的索引
                    'heading_level': str,      # 'major' / 'sub'
                    'section_type': str,       # 匹配的标准章节类型
                    'category': str,           # 资料库分类
                    'match_score': float,      # 匹配置信度
                    'content_start': int,      # 内容起始段落索引（标题下一段）
                    'content_end': int,        # 内容结束段落索引（下一个标题前）
                    'has_table': bool,         # 该章节区域内是否包含表格
                }
            ],
            'total_paragraphs': int,
            'total_tables': int,
            'message': str,
        }
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        return {
            'success': False,
            'sections': [],
            'total_paragraphs': 0,
            'total_tables': 0,
            'message': f'无法打开格式模板文件: {e}',
        }

    paragraphs = doc.paragraphs
    total_paras = len(paragraphs)
    total_tables = len(doc.tables)

    if total_paras == 0:
        return {
            'success': False,
            'sections': [],
            'total_paragraphs': 0,
            'total_tables': total_tables,
            'message': '格式模板文件中没有段落内容',
        }

    # 第一遍扫描：识别所有章节标题及其位置
    raw_headings = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        matched = _match_heading(text, para)
        if matched:
            level, name, num = matched
            raw_headings.append({
                'para_index': i,
                'heading_text': text,
                'section_name': name.strip(),
                'heading_level': level,
                'order_num': num,
            })

    if not raw_headings:
        return {
            'success': False,
            'sections': [],
            'total_paragraphs': total_paras,
            'total_tables': total_tables,
            'message': '未能在格式模板中识别到章节标题，请检查文件格式',
        }

    # 构建表格位置索引（段落索引 → 表格存在）
    table_para_ranges = _build_table_index(doc)

    # 第二遍：计算每个章节的内容范围，分类章节类型
    sections = []
    for idx, heading in enumerate(raw_headings):
        content_start = heading['para_index'] + 1
        if idx + 1 < len(raw_headings):
            content_end = raw_headings[idx + 1]['para_index']
        else:
            content_end = total_paras

        # 检查内容范围内是否有表格
        has_table = _range_has_table(table_para_ranges, content_start, content_end)

        # 分类章节
        sec_type, category, score = classify_section(heading['section_name'])

        sections.append({
            'order': idx + 1,
            'section_name': heading['section_name'],
            'heading_text': heading['heading_text'],
            'para_index': heading['para_index'],
            'heading_level': heading['heading_level'],
            'section_type': sec_type,
            'category': category,
            'match_score': score,
            'content_start': content_start,
            'content_end': content_end,
            'has_table': has_table,
        })

    return {
        'success': True,
        'sections': sections,
        'total_paragraphs': total_paras,
        'total_tables': total_tables,
        'message': f'成功识别 {len(sections)} 个章节',
    }


def _match_heading(text, paragraph=None):
    """
    尝试将文本匹配为章节标题。

    Returns:
        (level, name, order_num) 或 None
        level: 'major' / 'sub'
        name: 清理后的章节名称
        order_num: 序号（整数）
    """
    # 跳过过长文本（不太可能是标题）
    if len(text) > 100:
        return None

    # 跳过明显的正文内容
    if text.startswith('注：') or text.startswith('备注') or text.startswith('说明'):
        return None

    for pattern, level in HEADING_PATTERNS:
        m = re.match(pattern, text)
        if m:
            num_str = m.group(1)
            name = m.group(2).strip()

            # 清理名称末尾的冒号和标点
            name = re.sub(r'[：:]+$', '', name).strip()

            # 排除太短或太长的名称
            if len(name) < 2 or len(name) > 60:
                continue

            # 转换序号
            if level in ('major', 'sub'):
                order_num = _cn_to_int(num_str)
            else:
                order_num = int(num_str) if num_str.isdigit() else 0

            # 映射 sub_num → sub, major_num → major
            if level == 'sub_num':
                level = 'sub'
            elif level == 'major_num':
                level = 'major'

            return level, name, order_num

    # 额外检查：使用 Word 标题样式的段落
    if paragraph and _is_heading_style(paragraph):
        name = text.strip()
        if 2 <= len(name) <= 60:
            return 'major', name, 0

    return None


def _build_table_index(doc):
    """
    构建文档中表格的位置索引。
    返回一个列表，每项为 (before_para_index, after_para_index) 表示表格大致位于哪些段落之间。

    注意：python-docx 中表格和段落是 body 的平级子元素，
    通过遍历 body 的 XML 子元素来确定表格相对于段落的位置。
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    table_ranges = []
    para_count = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_count += 1
        elif tag == 'tbl':
            # 表格出现在第 para_count 个段落之后
            table_ranges.append(para_count)

    return table_ranges


def _range_has_table(table_positions, content_start, content_end):
    """检查指定段落范围内是否包含表格"""
    for tbl_pos in table_positions:
        if content_start <= tbl_pos <= content_end:
            return True
    return False


def get_section_content_text(filepath, para_start, para_end):
    """
    获取格式模板中指定范围的段落文本（用于预览）。

    Args:
        filepath: 格式模板文件路径
        para_start: 起始段落索引
        para_end: 结束段落索引

    Returns:
        str: 段落文本，换行分隔
    """
    doc = Document(filepath)
    lines = []
    for i in range(para_start, min(para_end, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            lines.append(text)
    return '\n'.join(lines)
