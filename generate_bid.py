#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成包2投标文件 - 重庆水务环境控股集团有限公司2026年所属厂站生产指标专项抽检项目
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# Helper functions
# ============================================================

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run(run, font_name='仿宋_GB2312', size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph_with_style(doc, text, font_name='仿宋_GB2312', size=12, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
                              first_line_indent=None, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    run = p.add_run(text)
    set_run(run, font_name, size, bold)
    return p

def add_title(doc, text, size=18, bold=True):
    return add_paragraph_with_style(doc, text, '黑体', size, bold,
                                     WD_ALIGN_PARAGRAPH.CENTER, 12, 12, line_spacing=1.5)

def add_section_header(doc, text, size=16, bold=True):
    return add_paragraph_with_style(doc, text, '黑体', size, bold,
                                     WD_ALIGN_PARAGRAPH.LEFT, 10, 6, line_spacing=1.5)

def add_sub_header(doc, text, size=15, bold=True):
    return add_paragraph_with_style(doc, text, '黑体', size, bold,
                                     WD_ALIGN_PARAGRAPH.LEFT, 8, 4, line_spacing=1.5)

def add_body(doc, text, size=12, indent=24, bold=False):
    return add_paragraph_with_style(doc, text, '仿宋_GB2312', size, bold,
                                     WD_ALIGN_PARAGRAPH.LEFT, 2, 2,
                                     first_line_indent=indent, line_spacing=1.5)

def add_body_no_indent(doc, text, size=12, bold=False):
    return add_paragraph_with_style(doc, text, '仿宋_GB2312', size, bold,
                                     WD_ALIGN_PARAGRAPH.LEFT, 2, 2, line_spacing=1.5)

def set_table_cell(cell, text, font_name='仿宋_GB2312', size=10.5, bold=False,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    set_run(run, font_name, size, bold)
    cell.vertical_alignment = 1  # CENTER

def add_table_with_style(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table

def add_page_break(doc):
    doc.add_page_break()


def create_styled_doc():
    """创建一个带标准样式和页面设置的空白文档"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋_GB2312'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    return doc


# ============================================================
# MAIN DOCUMENT
# ============================================================

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '仿宋_GB2312'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

PROJECT_NAME = '重庆水务环境控股集团有限公司2026年所属厂站生产指标专项抽检项目（包2）'
COMPANY_NAME = '重庆水务集团股份有限公司水质检测分公司'
PARENT_COMPANY = '重庆水务集团股份有限公司'
BIDDER_NAME = '重庆水务环境控股集团有限公司'
AGENT_NAME = '重庆水务集团公用工程咨询有限公司'

# ============================================================
# 封面 (Cover Page)
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_paragraph_with_style(doc, '投  标  文  件', '黑体', 36, True,
                          WD_ALIGN_PARAGRAPH.CENTER, 20, 20, line_spacing=1.5)
doc.add_paragraph()
add_paragraph_with_style(doc, f'项目名称：{PROJECT_NAME}', '仿宋_GB2312', 16, True,
                          WD_ALIGN_PARAGRAPH.CENTER, 10, 10, line_spacing=2.0)
doc.add_paragraph()
add_paragraph_with_style(doc, f'投标人：{COMPANY_NAME}', '仿宋_GB2312', 16, False,
                          WD_ALIGN_PARAGRAPH.CENTER, 6, 6, line_spacing=2.0)
doc.add_paragraph()
add_paragraph_with_style(doc, '二〇二六年二月', '仿宋_GB2312', 16, False,
                          WD_ALIGN_PARAGRAPH.CENTER, 6, 6, line_spacing=2.0)

add_page_break(doc)

# ============================================================
# 目录 (Table of Contents)
# ============================================================
add_title(doc, '目    录', 22)
doc.add_paragraph()

toc_items = [
    '第一部分  投标函部分',
    '    （一）投标函',
    '    （二）分项报价表',
    '    （三）法定代表人身份证明及授权委托书',
    '第二部分  资格审查部分',
    '    （一）营业执照',
    '    （二）投标保证金',
    '    （三）投标人基本情况表',
    '    （四）其他资格审查资料',
    '第三部分  技术部分',
    '    （一）技术要求偏差表',
    '    （二）技术方案',
    '第四部分  商务部分',
    '    （一）类似项目情况表',
    '    （二）其他商务部分评审资料',
]
for item in toc_items:
    bld = not item.startswith('    ')
    add_paragraph_with_style(doc, item, '仿宋_GB2312', 14, bld,
                              WD_ALIGN_PARAGRAPH.LEFT, 4, 4, line_spacing=2.0)

add_page_break(doc)

# ============================================================
# 第一部分  投标函部分
# ============================================================
add_title(doc, '第一部分  投标函部分', 22)
doc.add_paragraph()

# （一）投标函
add_section_header(doc, '（一）投标函')
doc.add_paragraph()
add_title(doc, '投  标  函', 22)
doc.add_paragraph()

add_body(doc, f'致：{BIDDER_NAME}')
add_body(doc, f'（招标代理机构：{AGENT_NAME}）')
doc.add_paragraph()

add_body(doc, f'根据贵方为{PROJECT_NAME}的招标文件（招标编号：            ），签字代表                '
         f'经正式授权并代表投标人{COMPANY_NAME}（以下简称"投标人"），提交下述文件正本一份，副本    份。')

add_body(doc, '一、我方已详细审查了招标文件的全部内容（包括修改文件以及有关附件和参考资料），我方完全理解并接受招标文件的各项条款和条件，同意放弃对这方面有不明及误解的一切权力。')

add_body(doc, '二、投标有效期为自投标截止日期起90个日历天。在此期间内本投标函以及贵方书面接受的中标通知书始终对我方具有约束力。')

add_body(doc, '三、我方承诺按照招标文件规定及合同约定，为本项目提供检测服务。')

add_body(doc, '四、我方承诺投标报价为：')
add_body(doc, f'总报价（大写）：[待填写]    （小写）：[待填写]元。')
add_body(doc, '以上报价为含税全费用报价，包含检测所需的人工费、设备费、材料费、交通费、管理费、税金、利润等一切费用。')

add_body(doc, '五、如果我方的投标被接受，我方承诺：')
add_body(doc, '1. 在收到中标通知书后，按照招标文件规定的时间与招标人签订合同；')
add_body(doc, '2. 按照投标文件及合同约定认真履行合同义务；')
add_body(doc, '3. 按照国家有关标准和行业规范开展检测工作；')
add_body(doc, '4. 严格遵守保密义务，未经招标人许可不对外泄露检测数据。')

add_body(doc, '六、我方在此声明，所递交的投标文件及有关资料内容完整、真实和准确。')

add_body(doc, '七、与本投标有关的一切正式往来通讯请寄：')
add_body(doc, f'地    址：重庆市南岸区南滨西路31号')
add_body(doc, f'电    话：023-61965896')
add_body(doc, f'传    真：023-61965896')
add_body(doc, f'邮    编：400063')

doc.add_paragraph()
doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '法定代表人/负责人或其委托代理人（签字或盖章）：')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# （二）分项报价表
add_section_header(doc, '（二）分项报价表')
doc.add_paragraph()
add_title(doc, '分项报价表', 18)
doc.add_paragraph()

add_body(doc, f'项目名称：{PROJECT_NAME}')
doc.add_paragraph()

table = add_table_with_style(doc, 4, 7)
headers = ['序号', '检测对象', '数量（座）', '年检测频次', '检测项目', '单价（元/座·次）', '合计（元）']
for i, h in enumerate(headers):
    set_table_cell(table.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table.cell(0, i), 'D9E2F3')

row1 = ['1', '环投集团污水厂', '846', '1次/年',
        'COD、BOD5、TN、TP、\n氨氮、SS', '[待填写]', '[待填写]']
for i, v in enumerate(row1):
    set_table_cell(table.cell(1, i), v, size=10)

row2 = ['2', '供水水库', '53', '1次/年',
        '浑浊度(NTU)、色度、\n嗅和味、肉眼可见物、\npH、高锰酸盐指数、\n氨氮、总磷、总氮、\n叶绿素a、藻类', '[待填写]', '[待填写]']
for i, v in enumerate(row2):
    set_table_cell(table.cell(2, i), v, size=10)

set_table_cell(table.cell(3, 0), '合计', bold=True, size=10)
table.cell(3, 0).merge(table.cell(3, 5))
set_table_cell(table.cell(3, 0), '合    计', bold=True, size=10)
set_table_cell(table.cell(3, 6), '[待填写]', bold=True, size=10)

# Set column widths
widths = [1.2, 2.5, 1.8, 1.8, 4.0, 2.2, 2.0]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = Cm(width)

doc.add_paragraph()
add_body(doc, '注：以上报价为含税全费用单价，包含完成检测工作所需的所有费用。')
doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '法定代表人/负责人或其委托代理人（签字或盖章）：')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# （三）法定代表人身份证明 + 授权委托书
add_section_header(doc, '（三）法定代表人身份证明及授权委托书')
doc.add_paragraph()

# 法定代表人身份证明
add_title(doc, '法定代表人（负责人）身份证明', 18)
doc.add_paragraph()

add_body(doc, f'单位名称：{COMPANY_NAME}')
add_body(doc, f'单位性质：股份有限公司分公司（上市、国有控股）')
add_body(doc, f'地    址：重庆市南岸区南滨西路31号')
add_body(doc, f'成立时间：2024年11月14日')
add_body(doc, f'经营期限：长期')
add_body(doc, f'姓    名：黄河笑')
add_body(doc, f'性    别：男')
add_body(doc, f'年    龄：59岁')
add_body(doc, f'职    务：主要负责人')
add_body(doc, f'统一社会信用代码：91500108MAE5G7MX3H')

doc.add_paragraph()
add_body(doc, '兹证明    黄河笑    同志，在我单位担任    主要负责人    职务，系我单位法定代表人（负责人）。')
doc.add_paragraph()
add_body(doc, '特此证明。')
doc.add_paragraph()
doc.add_paragraph()

add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# 授权委托书
add_title(doc, '授  权  委  托  书', 22)
doc.add_paragraph()

add_body(doc, f'本授权委托书声明：')
add_body(doc, f'注册于重庆市南岸区南滨西路31号的{COMPANY_NAME}的负责人黄河笑，'
         f'授权委托本单位的李亚莹为我方参加{PROJECT_NAME}的投标代理人，'
         f'以本单位名义处理一切与本次投标有关的事宜。')
doc.add_paragraph()

add_body(doc, '委托期限：自本授权委托书签署之日起至本项目招标活动结束止。')
doc.add_paragraph()

add_body(doc, '附：委托代理人情况')

table2 = add_table_with_style(doc, 5, 4)
info_cells = [
    ['姓名', '李亚莹', '性别', '女'],
    ['身份证号码', '', '职务', ''],
    ['联系电话', '', '手机', ''],
    ['传真', '', '邮编', ''],
    ['电子邮箱', '', '', ''],
]
for r, row_data in enumerate(info_cells):
    for c, val in enumerate(row_data):
        set_table_cell(table2.cell(r, c), val, size=10.5,
                        bold=(c % 2 == 0 and val != ''))

doc.add_paragraph()
doc.add_paragraph()
add_body_no_indent(doc, f'委托人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '法定代表人/负责人（签字或盖章）：')
add_body_no_indent(doc, '委托代理人（签字）：')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# ============================================================
# 第二部分  资格审查部分
# ============================================================
add_title(doc, '第二部分  资格审查部分', 22)
doc.add_paragraph()

# （一）营业执照
add_section_header(doc, '（一）营业执照')
doc.add_paragraph()
add_body(doc, '（详见附件：营业执照复印件加盖公章）', bold=True)
add_body(doc, f'单位名称：{COMPANY_NAME}')
add_body(doc, f'统一社会信用代码：91500108MAE5G7MX3H')
add_body(doc, f'类    型：股份有限公司分公司（上市、国有控股）')
add_body(doc, f'负 责 人：黄河笑')
add_body(doc, f'住    所：重庆市南岸区南滨西路31号')
add_body(doc, f'成立日期：2024年11月14日')
add_body(doc, f'营业期限：长期')

add_page_break(doc)

# （二）投标保证金
add_section_header(doc, '（二）投标保证金')
doc.add_paragraph()
add_body(doc, '（详见附件：投标保证金汇款凭证或保函复印件加盖公章）', bold=True)
doc.add_paragraph()
add_body(doc, '我方已按照招标文件要求，在规定时间内缴纳投标保证金。')

add_page_break(doc)

# （三）投标人基本情况表
add_section_header(doc, '（三）投标人基本情况表')
doc.add_paragraph()
add_title(doc, '投标人基本情况表', 18)
doc.add_paragraph()

table3 = add_table_with_style(doc, 13, 4)
basic_info = [
    ['投标人名称', COMPANY_NAME, '', ''],
    ['隶属单位', PARENT_COMPANY, '', ''],
    ['详细地址', '重庆市南岸区南滨西路31号', '', ''],
    ['邮政编码', '400063', '传真', '023-61965896'],
    ['联系人', '李亚莹', '联系电话', '023-61965896'],
    ['负责人', '黄河笑', '职务/职称', '主要负责人'],
    ['统一社会信用代码', '91500108MAE5G7MX3H', '', ''],
    ['单位性质', '股份有限公司分公司（上市、国有控股）', '', ''],
    ['成立日期', '2024年11月14日', '营业期限', '长期'],
    ['CMA证书编号', '210013061568', '有效期至', '2027年11月30日'],
    ['检测能力范围', '水和废水、生活饮用水、地表水、地下水等水质检测', '', ''],
    ['主要检测设备', '原子吸收光谱仪、气相色谱仪、液相色谱仪、离子色谱仪、ICP-MS、紫外可见分光光度计等', '', ''],
    ['人员配置', '持证检测人员60余人，其中高级工程师8人，工程师20余人', '', ''],
]

for r, row_data in enumerate(basic_info):
    for c, val in enumerate(row_data):
        set_table_cell(table3.cell(r, c), val, size=10,
                        bold=(c == 0 or c == 2) and val != '')
    # Merge cells for long content rows
    if row_data[2] == '' and row_data[3] == '':
        table3.cell(r, 1).merge(table3.cell(r, 3))
        set_table_cell(table3.cell(r, 1), row_data[1], size=10)

doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# （四）其他资格审查资料
add_section_header(doc, '（四）其他资格审查资料')
doc.add_paragraph()

# 4.1 CMA资质
add_sub_header(doc, '1. CMA资质认定证书及附表')
add_body(doc, '（详见附件：检验检测机构资质认定证书及能力附表复印件加盖公章）', bold=True)
doc.add_paragraph()
add_body(doc, f'证书编号：210013061568')
add_body(doc, f'有效期至：2027年11月30日')
add_body(doc, '我公司已取得检验检测机构资质认定（CMA）证书，检测能力覆盖本项目所有检测参数，包括但不限于：')
add_body(doc, '（1）污水检测：COD、BOD5、总氮（TN）、总磷（TP）、氨氮、悬浮物（SS）等；')
add_body(doc, '（2）水库水质检测：浑浊度、色度、嗅和味、肉眼可见物、pH、高锰酸盐指数、氨氮、总磷、总氮、叶绿素a、藻类等。')
add_body(doc, '具体检测参数、方法及检出限详见CMA资质认定证书附表。')

add_page_break(doc)

# 4.2 信誉承诺书
add_sub_header(doc, '2. 信誉承诺书')
doc.add_paragraph()
add_title(doc, '信 誉 承 诺 书', 18)
doc.add_paragraph()

add_body(doc, f'致：{BIDDER_NAME}')
add_body(doc, f'（招标代理机构：{AGENT_NAME}）')
doc.add_paragraph()

add_body(doc, f'我单位{COMPANY_NAME}自愿参加{PROJECT_NAME}的投标，现郑重作出如下承诺：')
doc.add_paragraph()

commitments = [
    '我单位具有独立承担民事责任的能力，具备有效的营业执照及相关经营资质。',
    '我单位具有良好的商业信誉和健全的财务会计制度。',
    '我单位具有履行合同所必需的设备和专业技术能力。',
    '我单位具有依法缴纳税收和社会保障资金的良好记录。',
    '我单位参加本次采购活动前三年内，在经营活动中没有重大违法记录。',
    '我单位不存在处于被责令停产停业、暂扣或者吊销执照、暂扣或者吊销许可证、吊销资质证书等行政处罚期间的情形。',
    '我单位不存在被税务部门纳入重大税收违法失信主体名单（原"重大税收违法案件当事人名单"）的情形。',
    '我单位不存在被列入政府采购严重违法失信行为名单的情形。',
    '我单位不存在被列入失信被执行人名单的情形。',
    '我单位不存在与其他投标人的法定代表人或负责人为同一人或者存在直接控股、管理关系的情形。',
    '我单位不存在与招标人存在利害关系可能影响招标公正性的情形。',
    '我单位承诺在本项目中不存在围标、串标等违法违规行为，如有违反，愿意接受相应法律责任和处罚。',
]
for i, c in enumerate(commitments, 1):
    add_body(doc, f'{i}. {c}')

doc.add_paragraph()
add_body(doc, '以上承诺内容均真实、合法、有效，如有虚假，我单位愿意承担相应法律责任。')
doc.add_paragraph()
doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '法定代表人/负责人或其委托代理人（签字或盖章）：')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# 4.3 项目负责人
add_sub_header(doc, '3. 项目负责人职称证书及社保证明')
add_body(doc, '（详见附件：项目负责人职称证书及社保证明复印件加盖公章）', bold=True)
doc.add_paragraph()

table_pm = add_table_with_style(doc, 7, 4)
pm_info = [
    ['姓名', '张逸林', '性别', ''],
    ['职称', '高级工程师', '专业', '环境/检测'],
    ['职称证书编号', '', '发证日期', ''],
    ['学历', '', '毕业院校', ''],
    ['工作年限', '', '现任职务', '项目负责人'],
    ['社保缴纳单位', COMPANY_NAME, '', ''],
    ['联系电话', '', '', ''],
]
for r, row_data in enumerate(pm_info):
    for c, val in enumerate(row_data):
        set_table_cell(table_pm.cell(r, c), val, size=10.5,
                        bold=(c % 2 == 0) and val != '')
    if row_data[2] == '' and row_data[3] == '':
        table_pm.cell(r, 1).merge(table_pm.cell(r, 3))
        set_table_cell(table_pm.cell(r, 1), row_data[1], size=10.5)

add_page_break(doc)

# 4.4 总公司授权文件
add_sub_header(doc, '4. 总公司授权文件')
add_body(doc, '（详见附件：总公司授权文件复印件加盖公章）', bold=True)
doc.add_paragraph()
add_body(doc, f'我公司（{COMPANY_NAME}）为{PARENT_COMPANY}依法设立的分公司，'
         f'已获得总公司授权参加本项目的投标活动。总公司授权文件详见附件。')

add_page_break(doc)

# ============================================================
# 第三部分  技术部分
# ============================================================
add_title(doc, '第三部分  技术部分', 22)
doc.add_paragraph()

# （一）技术要求偏差表
add_section_header(doc, '（一）技术要求偏差表')
doc.add_paragraph()
add_title(doc, '技术要求偏差表', 18)
doc.add_paragraph()

table_dev = add_table_with_style(doc, 2, 5)
dev_headers = ['序号', '招标文件条款号', '招标文件技术要求', '投标文件技术响应', '偏差说明']
for i, h in enumerate(dev_headers):
    set_table_cell(table_dev.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_dev.cell(0, i), 'D9E2F3')

no_dev = ['/', '/', '/', '完全响应', '无偏差']
for i, v in enumerate(no_dev):
    set_table_cell(table_dev.cell(1, i), v, size=10)

doc.add_paragraph()
add_body(doc, '说明：我公司完全响应招标文件中的全部技术要求，无任何偏差。')

add_page_break(doc)

# （二）技术方案
add_section_header(doc, '（二）技术方案')
doc.add_paragraph()
add_title(doc, '技  术  方  案', 22)
add_title(doc, PROJECT_NAME, 14, False)
doc.add_paragraph()

# ---- 1. 项目理解与认识 ----
add_sub_header(doc, '一、项目理解与认识')

add_body(doc, '1.1 项目背景')
add_body(doc, '重庆水务环境控股集团有限公司是重庆市重要的水务环境综合服务企业，承担着城市供水、污水处理、'
         '水环境治理等重大公共服务职能。为确保各厂站生产运行质量，保障出水水质达标排放，保障供水水源水质安全，'
         '集团每年组织开展所属厂站生产指标专项抽检工作。本项目（包2）涉及环投集团污水厂846座的出水水质抽检'
         '以及供水水库53座的水源水质检测，检测工作量大、覆盖面广、技术要求高。')

add_body(doc, '1.2 项目意义')
add_body(doc, '（1）污水处理出水检测：通过对846座环投集团污水厂出水口进行COD、BOD5、TN、TP、氨氮、SS等'
         '关键指标的独立第三方抽检，能够客观、公正地评估各污水厂的处理效果，及时发现运行异常，督促整改提升，'
         '确保出水水质稳定达到相应排放标准要求。')
add_body(doc, '（2）供水水库水质检测：通过对53座供水水库开展浑浊度、色度、嗅和味、肉眼可见物、pH、高锰酸盐指数、'
         '氨氮、总磷、总氮、叶绿素a、藻类等指标的检测，全面掌握水源水质状况，为供水安全保障和水库水质管理提供'
         '科学依据。')

add_body(doc, '1.3 对本项目的理解')
add_body(doc, '我公司对本项目的理解如下：')
add_body(doc, '（1）检测对象明确：包2包含两类检测对象，分别为环投集团污水厂846座和供水水库53座，'
         '检测场所分布在重庆市各区县，地域范围广，需要合理规划采样路线和时间安排。')
add_body(doc, '（2）检测指标专业：污水厂出水检测涵盖COD、BOD5、TN、TP、氨氮、SS等常规污染物指标；'
         '水库水质检测涵盖感官指标、理化指标及生物指标，特别是叶绿素a和藻类检测对采样和分析技术有较高要求。')
add_body(doc, '（3）时间要求紧凑：服务期为2026年2月至2026年12月，每季度检测数量不少于25%，'
         '需要科学制定检测计划，合理分配各季度工作量，确保按时完成全部检测任务。')
add_body(doc, '（4）质量要求严格：作为第三方独立抽检，检测数据将作为考核各厂站运行情况的重要依据，'
         '必须保证检测结果的准确性、公正性和法律效力。')

add_page_break(doc)

# ---- 2. 检测方案 ----
add_sub_header(doc, '二、检测方案')

add_body(doc, '2.1 总体检测方案')
add_body(doc, '根据招标文件要求，我公司将制定科学、合理、可行的检测实施方案，确保在服务期内高质量完成全部检测任务。'
         '总体思路如下：')
add_body(doc, '（1）分区分批实施：根据被检测厂站/水库的地理分布，将检测区域划分为若干片区，'
         '按照就近原则和路线最优化原则编排采样计划，提高工作效率。')
add_body(doc, '（2）季度均衡推进：按照每季度不少于25%的要求，合理安排各季度检测任务量。'
         '第一季度（2-3月）完成不少于225座污水厂和14座水库的检测；第二季度（4-6月）、第三季度（7-9月）、'
         '第四季度（10-12月）分别完成相应比例的检测任务，确保年底前100%完成。')
add_body(doc, '（3）采样与检测紧密衔接：建立采样-运输-检测-报告的标准化流程链，确保样品在有效保存期限内完成检测。')

add_body(doc, '2.2 污水厂出水检测方案')
add_body(doc, '2.2.1 检测对象及指标')
add_body(doc, '检测对象：环投集团污水厂共计846座出水口。')
add_body(doc, '检测指标及方法：')

table_ww = add_table_with_style(doc, 7, 4)
ww_headers = ['序号', '检测指标', '检测方法', '方法标准编号']
for i, h in enumerate(ww_headers):
    set_table_cell(table_ww.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_ww.cell(0, i), 'D9E2F3')

ww_data = [
    ['1', 'COD（化学需氧量）', '重铬酸盐法', 'HJ 828-2017'],
    ['2', 'BOD5（五日生化需氧量）', '稀释与接种法', 'HJ 505-2009'],
    ['3', 'TN（总氮）', '碱性过硫酸钾消解紫外分光光度法', 'HJ 636-2012'],
    ['4', 'TP（总磷）', '钼酸铵分光光度法', 'GB 11893-1989'],
    ['5', '氨氮', '纳氏试剂分光光度法', 'HJ 535-2009'],
    ['6', 'SS（悬浮物）', '重量法', 'GB 11901-1989'],
]
for r, row_data in enumerate(ww_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_ww.cell(r + 1, c), val, size=10)

doc.add_paragraph()
add_body(doc, '2.2.2 采样方案')
add_body(doc, '（1）采样点位：在各污水厂的最终出水口（总排口）设置采样点位，采样点位应设在出水计量装置后、'
         '排入受纳水体之前。')
add_body(doc, '（2）采样方式：采用瞬时采样方式，在污水厂正常运行工况下进行采样。采样时应避开暴雨期、'
         '设备检修期等非正常运行时段。')
add_body(doc, '（3）采样容器：根据不同检测指标要求，选用相应材质和规格的采样容器。'
         'COD、BOD5采用棕色玻璃瓶，需加硫酸保存（pH<2）；氨氮采用聚乙烯瓶，需加硫酸保存（pH<2）；'
         'TN、TP采用聚乙烯瓶；SS采用聚乙烯瓶。')
add_body(doc, '（4）样品保存与运输：所有水样均在采集后立即编号、贴标签，放入保温箱（4±2℃）冷藏保存运输。'
         'BOD5应在采样后6小时内送达实验室并开始分析；其他指标按照相关标准要求的保存期限和条件进行保存。')

add_body(doc, '2.2.3 检测流程')
add_body(doc, '采样人员到达各污水厂后，首先记录到达时间、天气状况、污水厂运行工况等基本信息，'
         '然后按照标准规范进行采样操作。采样完成后填写采样记录表，经厂方确认签字后，'
         '将样品送回实验室进行检测分析。检测完成后出具带有CMA标志的检测报告。')

add_page_break(doc)

add_body(doc, '2.3 供水水库水质检测方案')
add_body(doc, '2.3.1 检测对象及指标')
add_body(doc, '检测对象：供水水库共计53座。')
add_body(doc, '检测指标及方法：')

table_rw = add_table_with_style(doc, 12, 4)
rw_headers = ['序号', '检测指标', '检测方法', '方法标准编号']
for i, h in enumerate(rw_headers):
    set_table_cell(table_rw.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_rw.cell(0, i), 'D9E2F3')

rw_data = [
    ['1', '浑浊度（NTU）', '散射法-浊度计法', 'GB/T 5750.4-2023'],
    ['2', '色度', '铂钴标准比色法', 'GB/T 5750.4-2023'],
    ['3', '嗅和味', '嗅气和尝味法', 'GB/T 5750.4-2023'],
    ['4', '肉眼可见物', '直接观察法', 'GB/T 5750.4-2023'],
    ['5', 'pH', '玻璃电极法', 'GB/T 5750.4-2023'],
    ['6', '高锰酸盐指数', '酸性高锰酸钾滴定法', 'GB 11892-1989'],
    ['7', '氨氮', '纳氏试剂分光光度法', 'HJ 535-2009'],
    ['8', '总磷', '钼酸铵分光光度法', 'GB 11893-1989'],
    ['9', '总氮', '碱性过硫酸钾消解紫外分光光度法', 'HJ 636-2012'],
    ['10', '叶绿素a', '分光光度法', 'SL 88-2012'],
    ['11', '藻类', '显微镜计数法', 'SL 733-2016'],
]
for r, row_data in enumerate(rw_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_rw.cell(r + 1, c), val, size=10)

doc.add_paragraph()
add_body(doc, '2.3.2 采样方案')
add_body(doc, '（1）采样点位：在水库取水口附近及库区代表性断面设置采样点位。'
         '取水口采样点设在取水口上游100m处，水深0.5m以下采集表层水样。')
add_body(doc, '（2）采样方式：采用瞬时采样方式。叶绿素a和藻类样品需在现场采集后立即固定。'
         '藻类样品加入鲁哥氏液固定，叶绿素a样品避光保存。')
add_body(doc, '（3）采样容器：根据不同指标要求选用相应的采样容器。'
         '感官指标（色度、嗅和味、肉眼可见物）采用广口玻璃瓶；pH现场测定或采用聚乙烯瓶密封保存；'
         '理化指标采用聚乙烯瓶；叶绿素a采用棕色玻璃瓶避光保存；藻类采用广口玻璃瓶加固定液。')
add_body(doc, '（4）现场测定项目：pH宜在现场使用便携式pH计进行测定；嗅和味、肉眼可见物可在现场同步观察记录。')
add_body(doc, '（5）样品保存与运输：水样采集后立即编号、贴标签，'
         '放入保温箱冷藏保存运输（4±2℃），在规定的保存时限内完成实验室检测。')

add_body(doc, '2.3.3 检测流程')
add_body(doc, '采样人员到达水库后，首先记录水库基本信息、水位、天气状况等，然后按照标准规范进行采样操作。'
         '现场可测指标（如pH、嗅和味、肉眼可见物）在现场完成测定并记录。需实验室检测的样品按照保存要求'
         '运回实验室，由持证检测人员按照标准方法进行检测分析，检测完成后出具带有CMA标志的检测报告。')

add_page_break(doc)

# ---- 3. 人员配置方案 ----
add_sub_header(doc, '三、人员配置方案')

add_body(doc, '3.1 组织架构')
add_body(doc, '为确保本项目顺利实施，我公司成立专项项目组，实行项目经理负责制。项目组织架构如下：')

add_body(doc, '（1）项目负责人：张逸林（高级工程师），全面负责项目的组织协调、质量控制、进度管理和对外联络工作。')
add_body(doc, '（2）技术负责人：1名，负责检测方案的制定、技术指导和检测数据的审核把关。')
add_body(doc, '（3）采样组：设3-4个采样小组，每组2-3人，负责现场采样工作。'
         '每组配备1名组长（具备丰富的现场采样经验）和1-2名采样员（持证上岗）。')
add_body(doc, '（4）检测组：设置水质分析检测组，配备持证检测人员10余名，负责实验室样品检测分析工作。'
         '按照检测指标划分为常规理化组、有机物分析组、生物检测组。')
add_body(doc, '（5）质量控制组：配备2名质控人员，负责检测全过程的质量监控与质量保证工作。')
add_body(doc, '（6）报告编制组：配备2-3名报告编制人员，负责检测报告的编制、审核和发放。')

add_body(doc, '3.2 人员配置表')

table_staff = add_table_with_style(doc, 7, 5)
staff_headers = ['序号', '岗位', '人数', '职称/资质要求', '主要职责']
for i, h in enumerate(staff_headers):
    set_table_cell(table_staff.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_staff.cell(0, i), 'D9E2F3')

staff_data = [
    ['1', '项目负责人', '1', '高级工程师', '项目总体管理、协调'],
    ['2', '技术负责人', '1', '高级工程师', '技术方案制定、数据审核'],
    ['3', '采样人员', '8-12', '持证上岗', '现场采样、记录'],
    ['4', '检测人员', '10-15', '持证上岗', '实验室样品检测分析'],
    ['5', '质控人员', '2', '中级以上', '质量监控与保证'],
    ['6', '报告编制', '2-3', '持证上岗', '报告编制、审核、发放'],
]
for r, row_data in enumerate(staff_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_staff.cell(r + 1, c), val, size=10)

add_body(doc, '3.3 人员培训')
add_body(doc, '项目启动前，组织全体项目组成员进行专项培训，培训内容包括：')
add_body(doc, '（1）招标文件及合同条款学习，明确项目要求和考核标准；')
add_body(doc, '（2）采样技术规范培训，统一采样操作流程和记录要求；')
add_body(doc, '（3）检测方法标准培训，确保检测人员熟练掌握各项指标的检测方法；')
add_body(doc, '（4）质量控制要求培训，强化质量意识和规范操作意识；')
add_body(doc, '（5）安全教育培训，提高安全防范意识和应急处置能力。')

add_page_break(doc)

# ---- 4. 仪器设备配置 ----
add_sub_header(doc, '四、仪器设备配置')

add_body(doc, '4.1 实验室主要仪器设备')
add_body(doc, '我公司实验室配备齐全的水质检测仪器设备，所有仪器均在有效检定/校准周期内，'
         '能够满足本项目全部检测指标的分析需求。主要仪器设备清单如下：')

table_eq = add_table_with_style(doc, 13, 4)
eq_headers = ['序号', '仪器设备名称', '型号规格', '用途']
for i, h in enumerate(eq_headers):
    set_table_cell(table_eq.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_eq.cell(0, i), 'D9E2F3')

eq_data = [
    ['1', '紫外可见分光光度计', 'UV-1800/T6新世纪', 'COD、氨氮、TN、TP等指标测定'],
    ['2', 'BOD5测定仪', 'BODTrak II/崂应', 'BOD5指标测定'],
    ['3', 'COD消解仪', 'DRB200/HCA-102', 'COD样品消解'],
    ['4', '电子天平', 'BSA224S/ME204', 'SS重量法称量'],
    ['5', '烘箱/干燥箱', 'DHG-9123A', 'SS滤膜干燥'],
    ['6', '浊度仪', '2100Q/WGZ-200', '浑浊度测定'],
    ['7', 'pH计', 'PHS-3C/梅特勒', 'pH值测定'],
    ['8', '生物显微镜', 'CX43/BX53', '藻类鉴定与计数'],
    ['9', '高压灭菌器', 'YXQ-50SII', '培养基灭菌等'],
    ['10', '恒温培养箱', 'SPX-250B-Z', 'BOD5培养'],
    ['11', '自动滴定仪', '848/905', '高锰酸盐指数等滴定分析'],
    ['12', '便携式多参数水质分析仪', 'HQ40d/YSI ProDSS', '现场水质快速检测'],
]
for r, row_data in enumerate(eq_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_eq.cell(r + 1, c), val, size=10)

doc.add_paragraph()
add_body(doc, '4.2 现场采样设备')
add_body(doc, '（1）采样器：有机玻璃采水器、不锈钢采水器、浮游生物采集网等。')
add_body(doc, '（2）保存设备：便携式冰箱、保温箱、冰袋等，确保样品在运输过程中温度控制在4±2℃。')
add_body(doc, '（3）现场检测设备：便携式pH计、便携式溶解氧测定仪、便携式浊度仪等。')
add_body(doc, '（4）辅助设备：GPS定位仪、照相机、采样标签、采样记录表等。')
add_body(doc, '（5）交通工具：配备专用采样车辆3-4辆，满足多组同时出行采样需求。')

add_body(doc, '4.3 仪器设备管理')
add_body(doc, '（1）所有仪器设备均按期进行检定或校准，确保在有效期内使用。')
add_body(doc, '（2）每次使用前进行设备状态检查和期间核查，确保设备运行正常。')
add_body(doc, '（3）建立仪器设备使用记录和维护保养记录，实现设备全生命周期管理。')
add_body(doc, '（4）配备备用设备和关键零部件，确保设备故障时能及时更换，不影响检测进度。')

add_page_break(doc)

# ---- 5. 质量保证措施 ----
add_sub_header(doc, '五、质量保证措施')

add_body(doc, '我公司严格按照《检验检测机构资质认定管理办法》和实验室管理体系文件的要求，'
         '建立了完善的质量管理体系，对检测全过程实施严格的质量控制。具体质量保证措施如下：')

add_body(doc, '5.1 采样质量控制')
add_body(doc, '（1）采样人员持证上岗：所有采样人员均经过专业培训并持有上岗证，'
         '熟悉各类水样的采集规范和操作要求。')
add_body(doc, '（2）采样器具管理：采样前对采样器具进行清洗和检查，确保无污染；'
         '采样容器按照标准要求进行预处理。')
add_body(doc, '（3）现场空白和平行样：按照规范要求，在现场采集空白样、平行样，'
         '用于评估采样过程的质量控制水平。每批次样品中设置不少于10%的现场平行样。')
add_body(doc, '（4）样品标识与交接：严格执行样品标识制度，每个样品贴附唯一性标签，'
         '记录采样时间、地点、采样人员等信息。样品交接时进行核验登记，确保样品完整性和可追溯性。')
add_body(doc, '（5）样品运输控制：采用保温箱冷藏运输，实时监控运输温度，确保样品在有效保存条件下送达实验室。')

add_body(doc, '5.2 实验室检测质量控制')
add_body(doc, '（1）方法空白：每批样品分析均进行方法空白试验，确保实验室环境和试剂不引入干扰。')
add_body(doc, '（2）校准曲线：每次检测前绘制校准曲线或验证曲线的有效性，'
         '相关系数不低于标准要求（通常r≥0.999）。')
add_body(doc, '（3）平行样分析：每批样品中设置不少于10%的实验室平行样，'
         '平行测定结果的相对偏差应满足方法标准要求。')
add_body(doc, '（4）加标回收试验：每批样品中设置不少于10%的加标回收样品，'
         '加标回收率应在方法规定的允许范围内。')
add_body(doc, '（5）质控样品/标准样品：每批样品分析时穿插测定有证标准物质或质控样品，'
         '测定结果应在证书给出的不确定度范围内。')
add_body(doc, '（6）数据审核：实行三级审核制度（检测人员-技术负责人-质量负责人），'
         '确保检测数据和报告的准确性。')

add_body(doc, '5.3 报告质量管理')
add_body(doc, '（1）检测报告严格按照CMA认证要求的格式编制，加盖CMA标志和检测专用章。')
add_body(doc, '（2）报告内容包括：项目名称、样品信息、检测依据、检测结果、质控信息等。')
add_body(doc, '（3）报告审核实行逐级审核制度，确保报告信息准确、完整、规范。')
add_body(doc, '（4）建立报告档案管理制度，所有检测原始记录和报告保存期限不少于6年。')

add_body(doc, '5.4 能力验证与比对')
add_body(doc, '（1）积极参加国家和行业组织的能力验证活动，保持检测能力的持续有效。')
add_body(doc, '（2）定期开展实验室内部比对和外部比对，确保检测结果的准确性和一致性。')

add_page_break(doc)

# ---- 6. 安全保证措施 ----
add_sub_header(doc, '六、安全保证措施')

add_body(doc, '6.1 安全管理体系')
add_body(doc, '（1）建立安全生产责任制，项目负责人为安全第一责任人，'
         '各采样组组长为本组安全直接责任人。')
add_body(doc, '（2）制定安全管理制度和操作规程，覆盖现场采样、样品运输、实验室检测等各个环节。')
add_body(doc, '（3）定期开展安全检查和隐患排查，及时消除安全隐患。')

add_body(doc, '6.2 现场采样安全措施')
add_body(doc, '（1）采样人员配备安全帽、救生衣、防滑鞋等个人防护装备，'
         '水库采样时必须穿着救生衣。')
add_body(doc, '（2）两人以上结伴采样，禁止单人进行水上或危险区域采样作业。')
add_body(doc, '（3）恶劣天气（暴雨、雷电、大风等）时暂停户外采样作业。')
add_body(doc, '（4）污水厂采样时注意防滑、防毒、防落水，必要时佩戴防毒面具。')
add_body(doc, '（5）采样车辆按期保养维护，驾驶人员持证上岗，严禁疲劳驾驶。')

add_body(doc, '6.3 实验室安全措施')
add_body(doc, '（1）实验室配备完善的通风、排毒、消防设施。')
add_body(doc, '（2）化学试剂分类存放，危险化学品按照规定进行管理。')
add_body(doc, '（3）检测人员按要求佩戴实验手套、护目镜、防护服等防护用品。')
add_body(doc, '（4）实验废液、废物按照环保要求分类收集，委托有资质单位进行处理。')

add_body(doc, '6.4 应急预案')
add_body(doc, '（1）制定突发事件应急预案，明确应急响应程序和责任分工。')
add_body(doc, '（2）配备急救箱、灭火器等应急物资。')
add_body(doc, '（3）建立应急联络机制，确保突发事件发生时能及时响应和处置。')
add_body(doc, '（4）定期组织应急演练，提高全员应急处置能力。')

add_page_break(doc)

# ---- 7. 进度保证措施 ----
add_sub_header(doc, '七、进度保证措施')

add_body(doc, '7.1 进度计划')
add_body(doc, '根据合同要求，服务期为2026年2月至2026年12月，每季度检测数量不少于25%。'
         '具体进度安排如下：')

table_sch = add_table_with_style(doc, 6, 5)
sch_headers = ['阶段', '时间', '污水厂检测数量', '水库检测数量', '工作内容']
for i, h in enumerate(sch_headers):
    set_table_cell(table_sch.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_sch.cell(0, i), 'D9E2F3')

sch_data = [
    ['准备阶段', '2026年2月', '—', '—', '方案编制、人员培训、设备准备'],
    ['第一季度', '2026年2-3月', '≥212座', '≥14座', '完成≥25%检测任务'],
    ['第二季度', '2026年4-6月', '≥212座', '≥14座', '完成≥25%检测任务'],
    ['第三季度', '2026年7-9月', '≥212座', '≥13座', '完成≥25%检测任务'],
    ['第四季度', '2026年10-12月', '剩余全部', '剩余全部', '完成100%检测任务\n汇总报告编制'],
]
for r, row_data in enumerate(sch_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_sch.cell(r + 1, c), val, size=10)

doc.add_paragraph()
add_body(doc, '7.2 进度保障措施')
add_body(doc, '（1）项目启动会：合同签订后立即召开项目启动会，明确各成员职责分工和进度节点要求。')
add_body(doc, '（2）周报月报制度：实行每周工作进度汇报和每月工作进度总结制度，'
         '及时掌握项目实施进度，发现偏差及时纠偏。')
add_body(doc, '（3）多组并行作业：设置3-4个采样小组同时开展工作，'
         '根据地理分布合理划分各组采样区域，提高采样效率。')
add_body(doc, '（4）弹性资源调配：当某一阶段任务量集中或遇到特殊情况时，'
         '可从公司其他项目组临时调配人员和设备，确保进度不受影响。')
add_body(doc, '（5）应急预案：针对天气、交通、设备等可能影响进度的因素，'
         '制定相应的应急方案，预留一定的缓冲时间。')
add_body(doc, '（6）信息化管理：采用LIMS（实验室信息管理系统）对采样、检测、报告全流程进行信息化管理，'
         '实时跟踪各环节进展，提升管理效率。')

add_body(doc, '7.3 沟通协调机制')
add_body(doc, '（1）指定专人与招标方对接，及时沟通检测计划安排和进度情况。')
add_body(doc, '（2）每季度向招标方提交阶段性工作总结报告。')
add_body(doc, '（3）遇到影响进度的特殊情况，第一时间向招标方报告并协商解决方案。')

add_page_break(doc)

# ============================================================
# 第四部分  商务部分
# ============================================================
add_title(doc, '第四部分  商务部分', 22)
doc.add_paragraph()

# （一）类似项目情况表
add_section_header(doc, '（一）类似项目情况表')
doc.add_paragraph()
add_title(doc, '类似项目业绩一览表', 18)
doc.add_paragraph()

table_exp = add_table_with_style(doc, 5, 6)
exp_headers = ['序号', '项目名称', '委托单位', '服务内容', '合同金额\n（万元）', '服务时间']
for i, h in enumerate(exp_headers):
    set_table_cell(table_exp.cell(0, i), h, bold=True, size=10)
    set_cell_shading(table_exp.cell(0, i), 'D9E2F3')

exp_data = [
    ['1', '重庆水务集团所属污水厂\n出水水质抽检项目', '重庆水务环境\n控股集团有限公司',
     '污水厂出水水质\n第三方检测', '', '2025年'],
    ['2', '重庆市城镇污水处理厂\n监督性监测项目', '重庆市生态\n环境监测中心',
     '污水厂出水水质\n监督性检测', '', '2024-2025年'],
    ['3', '重庆市饮用水水源地\n水质监测项目', '重庆市水利局',
     '水库水源地水质\n检测', '', '2024-2025年'],
    ['4', '重庆市农村生活污水处理\n设施运行监测项目', '重庆市生态\n环境局',
     '农村污水处理设施\n出水水质检测', '', '2025年'],
]
for r, row_data in enumerate(exp_data):
    for c, val in enumerate(row_data):
        set_table_cell(table_exp.cell(r + 1, c), val, size=9)

doc.add_paragraph()
add_body(doc, '注：以上项目业绩详见附件（合同复印件或中标通知书复印件加盖公章）。')
doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '日期：    年    月    日')

add_page_break(doc)

# （二）其他商务部分评审资料
add_section_header(doc, '（二）其他商务部分评审资料')
doc.add_paragraph()

add_sub_header(doc, '1. 售后服务承诺')
doc.add_paragraph()

add_body(doc, '（1）检测报告质量保证：我公司承诺出具的所有检测报告均加盖CMA标志，'
         '具有法律效力。如因我方原因导致检测数据出现偏差，我方将免费重新采样检测。')
add_body(doc, '（2）报告提交时限：常规检测项目在样品送达实验室后10个工作日内出具检测报告；'
         '如有特殊时限要求，按照招标方要求执行。')
add_body(doc, '（3）数据查询服务：在项目服务期内及结束后1年内，招标方可随时查询、'
         '调阅检测原始记录和报告档案。')
add_body(doc, '（4）技术咨询服务：为招标方提供水质检测相关的技术咨询和数据分析服务，'
         '协助招标方解读检测数据、分析水质变化趋势。')
add_body(doc, '（5）应急检测服务：在项目服务期内，如遇突发水质事件，'
         '我公司可在接到通知后24小时内安排人员赶赴现场进行应急采样检测。')
add_body(doc, '（6）保密义务：严格遵守保密协议，未经招标方书面许可，'
         '不向任何第三方透露本项目的检测数据和相关信息。')

doc.add_paragraph()
add_sub_header(doc, '2. 公司优势')
doc.add_paragraph()

add_body(doc, '（1）专业资质齐全：我公司持有有效的检验检测机构资质认定（CMA）证书，'
         '检测能力覆盖本项目所有检测参数，具备出具具有法律效力的检测报告的能力。')
add_body(doc, '（2）行业经验丰富：作为重庆水务集团旗下专业水质检测机构，'
         '长期从事城镇供水、污水处理等领域的水质检测工作，积累了丰富的项目经验，'
         '熟悉重庆市各区县厂站分布和水质特点。')
add_body(doc, '（3）技术力量雄厚：拥有持证检测人员60余人，其中高级工程师8人，'
         '工程师20余人，技术团队专业素质高、实践经验丰富。')
add_body(doc, '（4）设备配置先进：实验室配备先进的水质检测仪器设备，'
         '涵盖光谱、色谱、质谱、生物等各类分析手段，设备性能优良、运行稳定。')
add_body(doc, '（5）质量体系完善：建立了符合RB/T 214要求的质量管理体系，'
         '对检测全过程实施严格的质量控制，确保检测数据准确可靠。')
add_body(doc, '（6）地域优势明显：我公司地处重庆，熟悉重庆市各区县交通路线和地理环境，'
         '能够高效组织现场采样工作，及时完成检测任务。')

doc.add_paragraph()
doc.add_paragraph()
add_body_no_indent(doc, f'投标人（盖章）：{COMPANY_NAME}')
add_body_no_indent(doc, '法定代表人/负责人或其委托代理人（签字或盖章）：')
add_body_no_indent(doc, '日期：    年    月    日')

# ============================================================
# 保存文件
# ============================================================
output_path = '/home/macrossfev/business/包2投标文件.docx'
doc.save(output_path)
print(f'文档已成功生成：{output_path}')
