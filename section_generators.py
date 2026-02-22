#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投标文件章节生成器
Generates formatted content for each bid section type, pulling data from the database.
"""

from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# Helper / formatting functions
# ============================================================

def set_cell_shading(cell, color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run(run, font_name='仿宋_GB2312', size=12, bold=False, color=None):
    """Configure a Run object with font, size, bold, and optional colour."""
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_with_style(doc, text, font_name='仿宋_GB2312', size=12,
                             bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=0, space_after=0,
                             first_line_indent=None, line_spacing=1.5):
    """Add a fully-styled paragraph to *doc* and return it."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    run = p.add_run(text)
    set_run(run, font_name, size, bold)
    return p


def add_title(doc, text, size=18, bold=True):
    """Centred heading in 黑体."""
    return add_paragraph_with_style(
        doc, text, '黑体', size, bold,
        WD_ALIGN_PARAGRAPH.CENTER, 12, 12, line_spacing=1.5)


def add_section_header(doc, text, size=16, bold=True):
    """Left-aligned section heading in 黑体."""
    return add_paragraph_with_style(
        doc, text, '黑体', size, bold,
        WD_ALIGN_PARAGRAPH.LEFT, 10, 6, line_spacing=1.5)


def add_sub_header(doc, text, size=15, bold=True):
    """Left-aligned sub-heading in 黑体."""
    return add_paragraph_with_style(
        doc, text, '黑体', size, bold,
        WD_ALIGN_PARAGRAPH.LEFT, 8, 4, line_spacing=1.5)


def add_body(doc, text, size=12, indent=24, bold=False):
    """Body paragraph in 仿宋_GB2312 with first-line indent."""
    return add_paragraph_with_style(
        doc, text, '仿宋_GB2312', size, bold,
        WD_ALIGN_PARAGRAPH.LEFT, 2, 2,
        first_line_indent=indent, line_spacing=1.5)


def add_body_no_indent(doc, text, size=12, bold=False):
    """Body paragraph in 仿宋_GB2312 without indent."""
    return add_paragraph_with_style(
        doc, text, '仿宋_GB2312', size, bold,
        WD_ALIGN_PARAGRAPH.LEFT, 2, 2, line_spacing=1.5)


def set_table_cell(cell, text, font_name='仿宋_GB2312', size=10.5,
                   bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Write *text* into a table cell with the given formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    set_run(run, font_name, size, bold)
    cell.vertical_alignment = 1  # CENTER


def add_table_with_style(doc, rows, cols):
    """Insert a centred Table Grid table and return it."""
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table


def add_page_break(doc):
    doc.add_page_break()


def _seal_and_date(doc, company_name):
    """Append the standard company-seal block and date line."""
    doc.add_paragraph()
    doc.add_paragraph()
    add_body_no_indent(doc, f'投标人（盖章）：{company_name}')
    add_body_no_indent(doc, '法定代表人/负责人或其委托代理人（签字或盖章）：')
    add_body_no_indent(doc, '日期：    年    月    日')


def _safe(value, placeholder='[待填写]'):
    """Return *value* if truthy, otherwise the placeholder string."""
    return value if value else placeholder


def create_styled_doc():
    """Create a blank Document with standard page setup and default font."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋_GB2312'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    return doc


# ============================================================
# Main dispatcher
# ============================================================

def generate_section(doc, section, bid, app_config):
    """Call the correct generator for *section.section_type*.

    Parameters
    ----------
    doc : Document
        The python-docx Document to append content to.
    section : BidSection
        The section row from the database.
    bid : BidProject
        The parent bid project.
    app_config : dict-like
        Flask app.config (provides UPLOAD_FOLDER, etc.).
    """
    generators = {
        '投标函': _gen_bid_letter,
        '报价表': _gen_pricing_table,
        '法定代表人证明': _gen_legal_rep_cert,
        '授权委托书': _gen_power_of_attorney,
        '营业执照': _gen_business_license,
        '投标保证金': _gen_bid_guarantee,
        '基本情况表': _gen_basic_info_table,
        'CMA证书': _gen_cma_cert,
        '信誉承诺书': _gen_credit_commitment,
        '人员资料': _gen_personnel_info,
        '技术方案': _gen_technical_proposal,
        '技术偏差表': _gen_technical_deviation,
        '业绩证明': _gen_performance_proof,
        '售后服务': _gen_after_sales,
    }
    gen_func = generators.get(section.section_type, _gen_other)
    gen_func(doc, section, bid, app_config)


# ============================================================
# Individual section generators
# ============================================================

def _gen_bid_letter(doc, section, bid, app_config):
    """投标函 - Formal bid letter."""
    project_name = _safe(bid.project_name)
    bidder_name = _safe(bid.bidder_name)
    agent_name = _safe(bid.agent_name)

    # Try to determine company name from BidPersonnel or notes
    company_name = _get_company_name(bid)

    add_title(doc, '投  标  函', 22)
    doc.add_paragraph()

    add_body(doc, f'致：{bidder_name}')
    add_body(doc, f'（招标代理机构：{agent_name}）')
    doc.add_paragraph()

    add_body(doc, f'根据贵方为{project_name}的招标文件（招标编号：            ），签字代表                '
             f'经正式授权并代表投标人{company_name}（以下简称"投标人"），提交下述文件正本一份，副本    份。')

    add_body(doc, '一、我方已详细审查了招标文件的全部内容（包括修改文件以及有关附件和参考资料），'
             '我方完全理解并接受招标文件的各项条款和条件，同意放弃对这方面有不明及误解的一切权力。')

    add_body(doc, '二、投标有效期为自投标截止日期起90个日历天。在此期间内本投标函以及贵方书面接受的'
             '中标通知书始终对我方具有约束力。')

    add_body(doc, '三、我方承诺按照招标文件规定及合同约定，为本项目提供检测服务。')

    add_body(doc, '四、我方承诺投标报价为：')
    add_body(doc, '总报价（大写）：[待填写]    （小写）：[待填写]元。')
    add_body(doc, '以上报价为含税全费用报价，包含检测所需的人工费、设备费、材料费、交通费、'
             '管理费、税金、利润等一切费用。')

    add_body(doc, '五、如果我方的投标被接受，我方承诺：')
    add_body(doc, '1. 在收到中标通知书后，按照招标文件规定的时间与招标人签订合同；')
    add_body(doc, '2. 按照投标文件及合同约定认真履行合同义务；')
    add_body(doc, '3. 按照国家有关标准和行业规范开展检测工作；')
    add_body(doc, '4. 严格遵守保密义务，未经招标人许可不对外泄露检测数据。')

    add_body(doc, '六、我方在此声明，所递交的投标文件及有关资料内容完整、真实和准确。')

    add_body(doc, '七、与本投标有关的一切正式往来通讯请寄：')
    add_body(doc, '地    址：[待填写]')
    add_body(doc, '电    话：[待填写]')
    add_body(doc, '传    真：[待填写]')
    add_body(doc, '邮    编：[待填写]')

    _seal_and_date(doc, company_name)


def _gen_pricing_table(doc, section, bid, app_config):
    """报价表 - Pricing / quotation table."""
    project_name = _safe(bid.project_name)
    company_name = _get_company_name(bid)

    add_title(doc, '分项报价表', 18)
    doc.add_paragraph()

    add_body(doc, f'项目名称：{project_name}')
    doc.add_paragraph()

    table = add_table_with_style(doc, 3, 5)
    headers = ['序号', '检测对象/服务内容', '数量', '单价', '合计']
    for i, h in enumerate(headers):
        set_table_cell(table.cell(0, i), h, bold=True, size=10)
        set_cell_shading(table.cell(0, i), 'D9E2F3')

    # Placeholder data row
    placeholder_row = ['1', '[待填写]', '[待填写]', '[待填写]', '[待填写]']
    for i, v in enumerate(placeholder_row):
        set_table_cell(table.cell(1, i), v, size=10)

    # Totals row
    set_table_cell(table.cell(2, 0), '合计', bold=True, size=10)
    table.cell(2, 0).merge(table.cell(2, 3))
    set_table_cell(table.cell(2, 0), '合    计', bold=True, size=10)
    set_table_cell(table.cell(2, 4), '[待填写]', bold=True, size=10)

    doc.add_paragraph()
    add_body(doc, '注：以上报价为含税全费用单价，包含完成本项目工作所需的所有费用。')
    _seal_and_date(doc, company_name)


def _gen_legal_rep_cert(doc, section, bid, app_config):
    """法定代表人证明 - Legal representative identity certificate."""
    from models import db, Personnel

    company_name = _get_company_name(bid)

    add_title(doc, '法定代表人（负责人）身份证明', 18)
    doc.add_paragraph()

    # Look up the legal representative from the Personnel table
    rep = Personnel.query.filter(
        (Personnel.position.contains('负责人')) |
        (Personnel.position.contains('法定代表人'))
    ).first()

    rep_name = rep.name if rep else '[待填写]'
    rep_gender = rep.gender if rep and rep.gender else '[待填写]'
    rep_position = rep.position if rep and rep.position else '主要负责人'

    add_body(doc, f'单位名称：{company_name}')
    add_body(doc, '单位性质：[待填写]')
    add_body(doc, '地    址：[待填写]')
    add_body(doc, '成立时间：[待填写]')
    add_body(doc, '经营期限：[待填写]')
    add_body(doc, f'姓    名：{rep_name}')
    add_body(doc, f'性    别：{rep_gender}')
    add_body(doc, '年    龄：[待填写]')
    add_body(doc, f'职    务：{rep_position}')
    add_body(doc, '统一社会信用代码：[待填写]')

    doc.add_paragraph()
    add_body(doc, f'兹证明    {rep_name}    同志，在我单位担任    {rep_position}    '
             f'职务，系我单位法定代表人（负责人）。')
    doc.add_paragraph()
    add_body(doc, '特此证明。')

    doc.add_paragraph()
    doc.add_paragraph()
    add_body_no_indent(doc, f'投标人（盖章）：{company_name}')
    add_body_no_indent(doc, '日期：    年    月    日')


def _gen_power_of_attorney(doc, section, bid, app_config):
    """授权委托书 - Power of attorney."""
    from models import db, Personnel, BidPersonnel

    company_name = _get_company_name(bid)
    project_name = _safe(bid.project_name)

    add_title(doc, '授  权  委  托  书', 22)
    doc.add_paragraph()

    # Find legal representative
    rep = Personnel.query.filter(
        (Personnel.position.contains('负责人')) |
        (Personnel.position.contains('法定代表人'))
    ).first()
    rep_name = rep.name if rep else '[待填写]'
    rep_position = rep.position if rep and rep.position else '负责人'

    # Find delegated agent from BidPersonnel
    delegate = BidPersonnel.query.filter(
        BidPersonnel.bid_project_id == bid.id,
        (BidPersonnel.role.contains('委托代理人')) |
        (BidPersonnel.role.contains('投标代理人'))
    ).first()

    delegate_name = '[待填写]'
    delegate_gender = '[待填写]'
    if delegate and delegate.personnel:
        delegate_name = delegate.personnel.name
        delegate_gender = delegate.personnel.gender or '[待填写]'

    add_body(doc, '本授权委托书声明：')
    add_body(doc, f'注册于[待填写]的{company_name}的{rep_position}{rep_name}，'
             f'授权委托本单位的{delegate_name}为我方参加{project_name}的投标代理人，'
             f'以本单位名义处理一切与本次投标有关的事宜。')
    doc.add_paragraph()

    add_body(doc, '委托期限：自本授权委托书签署之日起至本项目招标活动结束止。')
    doc.add_paragraph()

    add_body(doc, '附：委托代理人情况')

    table = add_table_with_style(doc, 5, 4)
    info_cells = [
        ['姓名', delegate_name, '性别', delegate_gender],
        ['身份证号码', '[待填写]', '职务', '[待填写]'],
        ['联系电话', '[待填写]', '手机', '[待填写]'],
        ['传真', '[待填写]', '邮编', '[待填写]'],
        ['电子邮箱', '[待填写]', '', ''],
    ]
    for r, row_data in enumerate(info_cells):
        for c, val in enumerate(row_data):
            set_table_cell(table.cell(r, c), val, size=10.5,
                           bold=(c % 2 == 0 and val != ''))

    doc.add_paragraph()
    doc.add_paragraph()
    add_body_no_indent(doc, f'委托人（盖章）：{company_name}')
    add_body_no_indent(doc, f'法定代表人/负责人（签字或盖章）：')
    add_body_no_indent(doc, '委托代理人（签字）：')
    add_body_no_indent(doc, '日期：    年    月    日')


def _gen_business_license(doc, section, bid, app_config):
    """营业执照 - Business licence reference."""
    company_name = _get_company_name(bid)

    add_section_header(doc, '营业执照')
    doc.add_paragraph()

    if section.attachment:
        add_body(doc, '（详见附件：营业执照复印件加盖公章）', bold=True)
        if section.attachment.notes:
            doc.add_paragraph()
            for line in section.attachment.notes.split('\n'):
                line = line.strip()
                if line:
                    add_body(doc, line)
    else:
        add_body(doc, '（详见附件：营业执照复印件加盖公章）', bold=True)

    doc.add_paragraph()
    add_body(doc, f'单位名称：{company_name}')
    add_body(doc, '统一社会信用代码：[待填写]')
    add_body(doc, '类    型：[待填写]')
    add_body(doc, '负 责 人：[待填写]')
    add_body(doc, '住    所：[待填写]')
    add_body(doc, '成立日期：[待填写]')
    add_body(doc, '营业期限：[待填写]')


def _gen_bid_guarantee(doc, section, bid, app_config):
    """投标保证金 - Bid guarantee / deposit."""
    add_section_header(doc, '投标保证金')
    doc.add_paragraph()
    add_body(doc, '（详见附件：投标保证金汇款凭证或保函复印件加盖公章）', bold=True)
    doc.add_paragraph()
    add_body(doc, '我方已按照招标文件要求，在规定时间内缴纳投标保证金。')


def _gen_basic_info_table(doc, section, bid, app_config):
    """基本情况表 - Company basic information table."""
    company_name = _get_company_name(bid)
    project_name = _safe(bid.project_name)

    add_title(doc, '投标人基本情况表', 18)
    doc.add_paragraph()

    table = add_table_with_style(doc, 10, 4)
    basic_info = [
        ['投标人名称', company_name, '', ''],
        ['详细地址', '[待填写]', '', ''],
        ['邮政编码', '[待填写]', '传真', '[待填写]'],
        ['联系人', '[待填写]', '联系电话', '[待填写]'],
        ['负责人', '[待填写]', '职务/职称', '[待填写]'],
        ['统一社会信用代码', '[待填写]', '', ''],
        ['单位性质', '[待填写]', '', ''],
        ['成立日期', '[待填写]', '营业期限', '[待填写]'],
        ['检测能力范围', '[待填写]', '', ''],
        ['主要检测设备', '[待填写]', '', ''],
    ]

    for r, row_data in enumerate(basic_info):
        for c, val in enumerate(row_data):
            set_table_cell(table.cell(r, c), val, size=10,
                           bold=(c == 0 or c == 2) and val != '')
        # Merge cells for rows where columns 3-4 are empty
        if row_data[2] == '' and row_data[3] == '':
            table.cell(r, 1).merge(table.cell(r, 3))
            set_table_cell(table.cell(r, 1), row_data[1], size=10)

    doc.add_paragraph()
    add_body_no_indent(doc, f'投标人（盖章）：{company_name}')
    add_body_no_indent(doc, '日期：    年    月    日')


def _gen_cma_cert(doc, section, bid, app_config):
    """CMA证书 - CMA certificate reference."""
    add_section_header(doc, 'CMA资质认定证书及附表')
    doc.add_paragraph()

    if section.attachment:
        add_body(doc, '（详见附件：检验检测机构资质认定证书及能力附表复印件加盖公章）', bold=True)
        if section.attachment.notes:
            doc.add_paragraph()
            for line in section.attachment.notes.split('\n'):
                line = line.strip()
                if line:
                    add_body(doc, line)
    else:
        add_body(doc, '（详见附件：检验检测机构资质认定证书及能力附表复印件加盖公章）', bold=True)

    doc.add_paragraph()
    add_body(doc, '证书编号：[待填写]')
    add_body(doc, '有效期至：[待填写]')
    add_body(doc, '我公司已取得检验检测机构资质认定（CMA）证书，检测能力覆盖本项目所有检测参数。'
             '具体检测参数、方法及检出限详见CMA资质认定证书附表。')


def _gen_credit_commitment(doc, section, bid, app_config):
    """信誉承诺书 - Credit / integrity commitment letter."""
    company_name = _get_company_name(bid)
    project_name = _safe(bid.project_name)
    bidder_name = _safe(bid.bidder_name)
    agent_name = _safe(bid.agent_name)

    add_title(doc, '信 誉 承 诺 书', 18)
    doc.add_paragraph()

    add_body(doc, f'致：{bidder_name}')
    add_body(doc, f'（招标代理机构：{agent_name}）')
    doc.add_paragraph()

    add_body(doc, f'我单位{company_name}自愿参加{project_name}的投标，现郑重作出如下承诺：')
    doc.add_paragraph()

    commitments = [
        '我单位具有独立承担民事责任的能力，具备有效的营业执照及相关经营资质。',
        '我单位具有良好的商业信誉和健全的财务会计制度。',
        '我单位具有履行合同所必需的设备和专业技术能力。',
        '我单位具有依法缴纳税收和社会保障资金的良好记录。',
        '我单位参加本次采购活动前三年内，在经营活动中没有重大违法记录。',
        '我单位不存在处于被责令停产停业、暂扣或者吊销执照、暂扣或者吊销许可证、'
        '吊销资质证书等行政处罚期间的情形。',
        '我单位不存在被税务部门纳入重大税收违法失信主体名单'
        '（原"重大税收违法案件当事人名单"）的情形。',
        '我单位不存在被列入政府采购严重违法失信行为名单的情形。',
        '我单位不存在被列入失信被执行人名单的情形。',
        '我单位不存在与其他投标人的法定代表人或负责人为同一人或者存在直接控股、管理关系的情形。',
        '我单位不存在与招标人存在利害关系可能影响招标公正性的情形。',
        '我单位承诺在本项目中不存在围标、串标等违法违规行为，如有违反，'
        '愿意接受相应法律责任和处罚。',
    ]
    for i, c in enumerate(commitments, 1):
        add_body(doc, f'{i}. {c}')

    doc.add_paragraph()
    add_body(doc, '以上承诺内容均真实、合法、有效，如有虚假，我单位愿意承担相应法律责任。')
    _seal_and_date(doc, company_name)


def _gen_personnel_info(doc, section, bid, app_config):
    """人员资料 - Personnel listing table."""
    from models import db, BidPersonnel

    company_name = _get_company_name(bid)

    add_title(doc, '项目人员配置表', 18)
    doc.add_paragraph()

    personnel_list = BidPersonnel.query.filter_by(
        bid_project_id=bid.id
    ).all()

    row_count = max(len(personnel_list), 1) + 1  # header + at least 1 data row
    table = add_table_with_style(doc, row_count, 5)
    headers = ['序号', '姓名', '职称', '职务', '角色']
    for i, h in enumerate(headers):
        set_table_cell(table.cell(0, i), h, bold=True, size=10)
        set_cell_shading(table.cell(0, i), 'D9E2F3')

    if personnel_list:
        for idx, bp in enumerate(personnel_list, 1):
            p = bp.personnel
            set_table_cell(table.cell(idx, 0), str(idx), size=10)
            set_table_cell(table.cell(idx, 1), p.name if p else '[待填写]', size=10)
            set_table_cell(table.cell(idx, 2), p.title if p and p.title else '[待填写]', size=10)
            set_table_cell(table.cell(idx, 3), p.position if p and p.position else '[待填写]', size=10)
            set_table_cell(table.cell(idx, 4), bp.role if bp.role else '[待填写]', size=10)
    else:
        for i in range(5):
            set_table_cell(table.cell(1, i), '[待填写]', size=10)

    doc.add_paragraph()
    add_body(doc, '注：以上人员相关职称证书、资格证书详见附件复印件（加盖公章）。')
    doc.add_paragraph()
    add_body_no_indent(doc, f'投标人（盖章）：{company_name}')
    add_body_no_indent(doc, '日期：    年    月    日')


def _gen_technical_proposal(doc, section, bid, app_config):
    """技术方案 - Technical proposal (custom content or skeleton)."""
    project_name = _safe(bid.project_name)

    add_title(doc, '技  术  方  案', 22)
    add_title(doc, project_name, 14, False)
    doc.add_paragraph()

    if section.custom_content:
        # Use the manually-provided content, splitting by newlines into body paragraphs
        for line in section.custom_content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            else:
                add_body(doc, line)
        return

    # Generate a skeleton proposal with standard headings
    skeleton_sections = [
        ('一、项目理解与认识', [
            '1.1 项目背景',
            f'[请根据{project_name}的招标文件要求，阐述项目背景情况。]',
            '1.2 项目意义',
            '[请阐述本项目的重要意义和目标。]',
            '1.3 对本项目的理解',
            '[请阐述对本项目检测内容、范围、技术要求等方面的理解。]',
        ]),
        ('二、检测/服务方案', [
            '2.1 总体方案',
            '[请描述本项目的总体实施方案和工作思路。]',
            '2.2 检测指标及方法',
            '[请列出本项目涉及的检测指标、采用的检测方法和标准。]',
            '2.3 采样方案',
            '[请描述采样点位设置、采样方式、采样频次、样品保存与运输等内容。]',
            '2.4 检测流程',
            '[请描述从采样到出具报告的完整检测流程。]',
        ]),
        ('三、人员配置方案', [
            '3.1 组织架构',
            '[请描述项目组织架构及各岗位职责。]',
            '3.2 人员配置',
            '[请列出拟投入本项目的主要人员及其资质情况。]',
            '3.3 人员培训',
            '[请描述项目启动前及实施过程中的培训计划。]',
        ]),
        ('四、仪器设备配置', [
            '4.1 实验室主要仪器设备',
            '[请列出开展本项目检测所需的主要仪器设备。]',
            '4.2 现场采样设备',
            '[请列出现场采样所需的设备和工具。]',
            '4.3 仪器设备管理',
            '[请描述仪器设备的检定/校准和维护管理措施。]',
        ]),
        ('五、质量保证措施', [
            '5.1 采样质量控制',
            '[请描述现场采样环节的质量控制措施。]',
            '5.2 实验室检测质量控制',
            '[请描述实验室检测环节的质量控制措施，包括方法空白、平行样、加标回收等。]',
            '5.3 报告质量管理',
            '[请描述检测报告的编制、审核和质量保证措施。]',
        ]),
        ('六、安全保证措施', [
            '6.1 安全管理体系',
            '[请描述安全管理组织体系和制度建设。]',
            '6.2 现场采样安全措施',
            '[请描述现场采样过程中的安全防护措施。]',
            '6.3 实验室安全措施',
            '[请描述实验室操作安全管理措施。]',
            '6.4 应急预案',
            '[请描述突发事件的应急处置方案。]',
        ]),
        ('七、进度保证措施', [
            '7.1 进度计划',
            '[请制定详细的项目实施进度计划。]',
            '7.2 进度保障措施',
            '[请描述确保按期完成全部检测任务的保障措施。]',
            '7.3 沟通协调机制',
            '[请描述与招标方的沟通协调机制和报告制度。]',
        ]),
    ]

    for heading, paragraphs in skeleton_sections:
        add_sub_header(doc, heading)
        for text in paragraphs:
            add_body(doc, text)
        doc.add_paragraph()


def _gen_technical_deviation(doc, section, bid, app_config):
    """技术偏差表 - Technical deviation table."""
    add_title(doc, '技术要求偏差表', 18)
    doc.add_paragraph()

    table = add_table_with_style(doc, 2, 5)
    headers = ['序号', '招标文件条款号', '招标文件技术要求', '投标文件技术响应', '偏差说明']
    for i, h in enumerate(headers):
        set_table_cell(table.cell(0, i), h, bold=True, size=10)
        set_cell_shading(table.cell(0, i), 'D9E2F3')

    no_dev = ['/', '/', '/', '完全响应', '无偏差']
    for i, v in enumerate(no_dev):
        set_table_cell(table.cell(1, i), v, size=10)

    doc.add_paragraph()
    add_body(doc, '说明：我公司完全响应招标文件中的全部技术要求，无任何偏差。')


def _gen_performance_proof(doc, section, bid, app_config):
    """业绩证明 - Past performance / project experience table."""
    from models import db, BidPerformance

    company_name = _get_company_name(bid)

    add_title(doc, '类似项目业绩一览表', 18)
    doc.add_paragraph()

    perf_list = BidPerformance.query.filter_by(
        bid_project_id=bid.id
    ).all()

    row_count = max(len(perf_list), 1) + 1
    table = add_table_with_style(doc, row_count, 6)
    headers = ['序号', '项目名称', '委托单位', '服务内容', '合同金额\n（万元）', '服务时间']
    for i, h in enumerate(headers):
        set_table_cell(table.cell(0, i), h, bold=True, size=10)
        set_cell_shading(table.cell(0, i), 'D9E2F3')

    if perf_list:
        for idx, bp in enumerate(perf_list, 1):
            perf = bp.performance
            set_table_cell(table.cell(idx, 0), str(idx), size=9)
            set_table_cell(table.cell(idx, 1),
                           perf.project_name if perf else '[待填写]', size=9)
            set_table_cell(table.cell(idx, 2),
                           perf.client_name if perf and perf.client_name else '[待填写]', size=9)
            set_table_cell(table.cell(idx, 3),
                           perf.description if perf and perf.description else '[待填写]', size=9)
            amount_str = str(perf.contract_amount) if perf and perf.contract_amount else '[待填写]'
            set_table_cell(table.cell(idx, 4), amount_str, size=9)
            # Format service time
            time_str = _format_service_time(perf)
            set_table_cell(table.cell(idx, 5), time_str, size=9)
    else:
        for i in range(6):
            set_table_cell(table.cell(1, i), '[待填写]', size=9)

    doc.add_paragraph()
    add_body(doc, '注：以上项目业绩详见附件（合同复印件或中标通知书复印件加盖公章）。')
    doc.add_paragraph()
    add_body_no_indent(doc, f'投标人（盖章）：{company_name}')
    add_body_no_indent(doc, '日期：    年    月    日')


def _gen_after_sales(doc, section, bid, app_config):
    """售后服务 - After-sales service commitments."""
    add_title(doc, '售后服务承诺', 18)
    doc.add_paragraph()

    add_body(doc, '（1）检测报告质量保证：我公司承诺出具的所有检测报告均加盖CMA标志，'
             '具有法律效力。如因我方原因导致检测数据出现偏差，我方将免费重新采样检测。')
    add_body(doc, '（2）报告提交时限：常规检测项目在样品送达实验室后10个工作日内出具检测报告；'
             '如有特殊时限要求，按照招标方要求执行。')
    add_body(doc, '（3）数据查询服务：在项目服务期内及结束后1年内，招标方可随时查询、'
             '调阅检测原始记录和报告档案。')
    add_body(doc, '（4）技术咨询服务：为招标方提供水质检测相关的技术咨询和数据分析服务，'
             '协助招标方解读检测数据、分析水质变化趋势。')
    add_body(doc, '（5）应急检测服务：在项目服务期内，如遇突发水质事件，'
             '我公司可在接到通知后24小时内安排人员赶赴现场进行应急采样检测。')
    add_body(doc, '（6）保密义务：严格遵守保密协议，未经招标方书面许可，'
             '不向任何第三方透露本项目的检测数据和相关信息。')


def _gen_other(doc, section, bid, app_config):
    """其他 / fallback - Unrecognised section type."""
    section_name = section.section_name or section.section_type or '其他资料'
    add_section_header(doc, section_name)
    doc.add_paragraph()

    if section.attachment:
        att_name = section.attachment.name or section_name
        add_body(doc, f'（详见附件：{att_name}）', bold=True)
        if section.attachment.notes:
            doc.add_paragraph()
            for line in section.attachment.notes.split('\n'):
                line = line.strip()
                if line:
                    add_body(doc, line)
    elif section.custom_content:
        for line in section.custom_content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            else:
                add_body(doc, line)
    else:
        add_body(doc, f'[{section_name} - 待填写]')


# ============================================================
# Composite document generators
# ============================================================

def generate_full_bid(bid, app_config):
    """Create a complete bid Document with cover page, TOC, and all sections.

    Returns
    -------
    doc : Document
        The assembled python-docx Document (not yet saved).
    """
    doc = create_styled_doc()

    company_name = _get_company_name(bid)
    project_name = _safe(bid.project_name)

    # ---- Cover page ----
    for _ in range(4):
        doc.add_paragraph()

    add_paragraph_with_style(
        doc, '投  标  文  件', '黑体', 36, True,
        WD_ALIGN_PARAGRAPH.CENTER, 20, 20, line_spacing=1.5)
    doc.add_paragraph()
    add_paragraph_with_style(
        doc, f'项目名称：{project_name}', '仿宋_GB2312', 16, True,
        WD_ALIGN_PARAGRAPH.CENTER, 10, 10, line_spacing=2.0)
    doc.add_paragraph()
    add_paragraph_with_style(
        doc, f'投标人：{company_name}', '仿宋_GB2312', 16, False,
        WD_ALIGN_PARAGRAPH.CENTER, 6, 6, line_spacing=2.0)
    doc.add_paragraph()

    # Date line - use Chinese year/month format
    now = datetime.now()
    cn_year = _arabic_to_cn_year(now.year)
    cn_month = _arabic_to_cn_month(now.month)
    add_paragraph_with_style(
        doc, f'{cn_year}年{cn_month}月', '仿宋_GB2312', 16, False,
        WD_ALIGN_PARAGRAPH.CENTER, 6, 6, line_spacing=2.0)

    add_page_break(doc)

    # ---- Table of contents ----
    add_title(doc, '目    录', 22)
    doc.add_paragraph()

    sections = sorted(bid.sections, key=lambda s: s.section_order or 0)
    for idx, sec in enumerate(sections, 1):
        display = f'    （{_cn_number(idx)}）{sec.section_name}'
        add_paragraph_with_style(
            doc, display, '仿宋_GB2312', 14, False,
            WD_ALIGN_PARAGRAPH.LEFT, 4, 4, line_spacing=2.0)

    add_page_break(doc)

    # ---- Sections ----
    for sec in sections:
        add_section_header(doc, sec.section_name)
        doc.add_paragraph()
        generate_section(doc, sec, bid, app_config)
        add_page_break(doc)

    return doc


def generate_section_preview(section, bid, app_config):
    """Create a Document containing only the given section's content.

    Returns
    -------
    doc : Document
        A single-section python-docx Document (not yet saved).
    """
    doc = create_styled_doc()
    add_section_header(doc, section.section_name or '[未命名章节]')
    doc.add_paragraph()
    generate_section(doc, section, bid, app_config)
    return doc


# ============================================================
# Internal utilities
# ============================================================

def _get_company_name(bid):
    """Derive the investing company name from bid data.

    Looks at BidPersonnel first (for a personnel record), then falls
    back to bid.notes, and finally to a placeholder.
    """
    from models import BidPersonnel
    try:
        bp = BidPersonnel.query.filter_by(bid_project_id=bid.id).first()
        if bp and bp.personnel and bp.personnel.social_security_unit:
            return bp.personnel.social_security_unit
    except Exception:
        pass

    if bid.notes:
        return bid.notes

    return '[投标单位名称待填写]'


def _format_service_time(perf):
    """Format a Performance record's service dates into a readable string."""
    if not perf:
        return '[待填写]'
    parts = []
    if perf.service_start:
        parts.append(perf.service_start.strftime('%Y年%m月'))
    if perf.service_end:
        parts.append(perf.service_end.strftime('%Y年%m月'))
    if parts:
        return ' - '.join(parts)
    return '[待填写]'


def _arabic_to_cn_year(year):
    """Convert an Arabic year (e.g. 2026) to Chinese characters."""
    digit_map = {
        '0': '〇', '1': '一', '2': '二', '3': '三', '4': '四',
        '5': '五', '6': '六', '7': '七', '8': '八', '9': '九',
    }
    return ''.join(digit_map.get(d, d) for d in str(year))


def _arabic_to_cn_month(month):
    """Convert a month number (1-12) to Chinese characters."""
    cn = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五', 6: '六',
        7: '七', 8: '八', 9: '九', 10: '十', 11: '十一', 12: '十二',
    }
    return cn.get(month, str(month))


def _cn_number(n):
    """Convert small integers (1-20) to Chinese ordinal characters."""
    cn = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
        16: '十六', 17: '十七', 18: '十八', 19: '十九', 20: '二十',
    }
    return cn.get(n, str(n))
