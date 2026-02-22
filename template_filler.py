#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板填充引擎
打开格式模板 .docx 文件，在各章节对应位置填充资料库匹配到的内容，
严格保留原始文档的字体、排版、表格样式等格式。
"""

import re
import os
import copy
from datetime import datetime, date
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 常见占位符模式
# ============================================================

# [xxx] 或 【xxx】 形式的占位符
BRACKET_PLACEHOLDER = re.compile(r'[\[【]\s*([^】\]]+?)\s*[】\]]')

# 下划线占位符（连续3个以上下划线）
UNDERLINE_PLACEHOLDER = re.compile(r'_{3,}')

# 常见占位符关键词 → 数据字段映射
PLACEHOLDER_FIELD_MAP = {
    # 公司信息
    '公司名称': 'company_name',
    '投标人名称': 'company_name',
    '投标人': 'company_name',
    '供应商名称': 'company_name',
    '单位名称': 'company_name',
    # 项目信息
    '项目名称': 'project_name',
    '项目编号': 'project_code',
    '招标编号': 'project_code',
    # 招标人
    '招标人': 'bidder_name',
    '采购人': 'bidder_name',
    '委托单位': 'bidder_name',
    # 代理机构
    '代理机构': 'agent_name',
    '招标代理': 'agent_name',
    # 金额
    '投标总价': 'total_price',
    '投标报价': 'total_price',
    '报价': 'total_price',
    '最高限价': 'max_price',
    # 日期
    '日期': 'current_date',
    '年月日': 'current_date',
    # 人员
    '法定代表人': 'legal_rep',
    '项目负责人': 'project_leader',
    '技术负责人': 'tech_leader',
    # 服务期
    '服务期': 'service_period',
    '服务期限': 'service_period',
    '工期': 'service_period',
}


def _get_project_data(bid, app_config=None):
    """
    从 BidProject 和关联数据中构建填充数据字典。
    """
    data = {
        'company_name': '[待填写]',
        'project_name': bid.project_name or '[待填写]',
        'project_code': '[待填写]',
        'bidder_name': bid.bidder_name or '[待填写]',
        'agent_name': bid.agent_name or '[待填写]',
        'total_price': '[待填写]',
        'max_price': f'{bid.max_price}万元' if bid.max_price else '[待填写]',
        'current_date': datetime.now().strftime('%Y年%m月%d日'),
        'legal_rep': '[待填写]',
        'project_leader': '[待填写]',
        'tech_leader': '[待填写]',
        'service_period': bid.service_period or '[待填写]',
    }

    # 从 notes 或 bid_personnel 中提取公司名称
    if bid.notes:
        data['company_name'] = bid.notes
    elif bid.bid_personnel:
        for bp in bid.bid_personnel:
            if bp.personnel and bp.personnel.social_security_unit:
                data['company_name'] = bp.personnel.social_security_unit
                break

    # 提取关键人员名称
    for bp in (bid.bid_personnel or []):
        if bp.role == '项目负责人' and bp.personnel:
            data['project_leader'] = bp.personnel.name
        elif bp.role == '技术负责人' and bp.personnel:
            data['tech_leader'] = bp.personnel.name
        elif bp.role == '法定代表人' and bp.personnel:
            data['legal_rep'] = bp.personnel.name

    return data


def fill_template(format_file_path, sections_info, bid, app_config=None):
    """
    在格式模板上填充内容，生成投标文件初稿。

    Args:
        format_file_path: 格式模板 .docx 文件路径
        sections_info: 章节信息列表（来自 format_parser 或数据库 BidSection）
            每项包含: section_type, para_index, content_start, content_end,
                      attachment, custom_content, personnel, performances
        bid: BidProject 对象
        app_config: Flask app.config

    Returns:
        Document: 填充后的 python-docx Document 对象
    """
    doc = Document(format_file_path)
    project_data = _get_project_data(bid, app_config)

    # 第一步：全局占位符替换（替换整个文档中的通用占位符）
    _replace_global_placeholders(doc, project_data)

    # 第二步：按章节填充特定内容
    # 需要从后往前处理，以免插入内容后段落索引偏移
    sorted_sections = sorted(sections_info, key=lambda s: s.get('content_start', 0), reverse=True)

    for sec_info in sorted_sections:
        section_type = sec_info.get('section_type', '')
        content_start = sec_info.get('content_start', 0)
        content_end = sec_info.get('content_end', 0)

        # 根据章节类型填充
        if section_type == '人员资料':
            _fill_personnel_section(doc, sec_info, bid, project_data, content_start, content_end)
        elif section_type == '业绩证明':
            _fill_performance_section(doc, sec_info, bid, project_data, content_start, content_end)
        elif section_type in ('投标函', '报价表', '法定代表人证明', '授权委托书'):
            _fill_form_section(doc, sec_info, bid, project_data, content_start, content_end)
        elif section_type == '技术方案':
            _fill_text_section(doc, sec_info, bid, project_data, content_start, content_end)
        elif sec_info.get('custom_content'):
            _fill_custom_content(doc, sec_info, project_data, content_start, content_end)

    return doc


def _replace_global_placeholders(doc, project_data):
    """
    替换文档中所有段落和表格中的占位符。
    保留原始 run 的格式。
    """
    # 处理所有段落
    for para in doc.paragraphs:
        _replace_paragraph_placeholders(para, project_data)

    # 处理所有表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_paragraph_placeholders(para, project_data)


def _replace_paragraph_placeholders(paragraph, project_data):
    """
    替换段落中的占位符，保留原始格式。

    策略：
    1. 先尝试在单个 run 内替换 [xxx] 占位符
    2. 如果占位符跨 run，合并后替换再重建
    3. 替换下划线占位符
    """
    full_text = paragraph.text
    if not full_text:
        return

    # 检查是否有占位符
    has_bracket = BRACKET_PLACEHOLDER.search(full_text)
    has_underline = UNDERLINE_PLACEHOLDER.search(full_text)

    if not has_bracket and not has_underline:
        return

    # 尝试在单个 run 内替换
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue

        # 替换 [xxx] 占位符
        new_text = BRACKET_PLACEHOLDER.sub(
            lambda m: _resolve_placeholder(m.group(1), project_data), run_text
        )
        # 替换下划线占位符（仅当前后有占位符关键词上下文时）
        if new_text != run_text:
            run.text = new_text

    # 如果单 run 替换后仍有跨 run 的占位符，尝试合并替换
    full_text_after = paragraph.text
    if BRACKET_PLACEHOLDER.search(full_text_after):
        _replace_cross_run_placeholders(paragraph, project_data)


def _replace_cross_run_placeholders(paragraph, project_data):
    """
    处理占位符跨多个 run 的情况。
    策略：记录所有 run 的文本和格式，合并替换后重新分配。
    """
    runs = paragraph.runs
    if not runs:
        return

    # 收集所有 run 的信息
    run_info = []
    for run in runs:
        run_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'underline': run.underline,
        })

    # 合并全文
    full_text = ''.join(r['text'] for r in run_info)
    new_text = BRACKET_PLACEHOLDER.sub(
        lambda m: _resolve_placeholder(m.group(1), project_data), full_text
    )

    if new_text == full_text:
        return

    # 用第一个 run 的格式写入全部新文本
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''


def _resolve_placeholder(key, project_data):
    """将占位符关键词解析为实际值"""
    key = key.strip()
    # 直接匹配
    field = PLACEHOLDER_FIELD_MAP.get(key)
    if field and field in project_data:
        return project_data[field]

    # 模糊匹配
    for ph_key, field_name in PLACEHOLDER_FIELD_MAP.items():
        if ph_key in key or key in ph_key:
            if field_name in project_data:
                return project_data[field_name]

    # 未识别的占位符保留原样
    return f'[{key}]'


def _fill_form_section(doc, sec_info, bid, project_data, content_start, content_end):
    """
    填充表单类章节（投标函、报价表、法定代表人证明、授权委托书等）。
    这些章节通常已有完整格式，只需替换占位符和填充表格。
    全局占位符替换已处理大部分内容，此处处理表格数据。
    """
    # 查找该章节范围内的表格并填充
    _fill_section_tables(doc, project_data, content_start, content_end)


def _fill_section_tables(doc, project_data, content_start, content_end):
    """填充指定段落范围内的表格单元格"""
    from docx.oxml.ns import qn as _qn

    body = doc.element.body
    para_count = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_count += 1
        elif tag == 'tbl':
            if content_start <= para_count <= content_end:
                # 这个表格在章节范围内，填充占位符
                from docx.table import Table
                tbl = Table(child, doc)
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_paragraph_placeholders(para, project_data)


def _fill_personnel_section(doc, sec_info, bid, project_data, content_start, content_end):
    """
    填充人员资料章节。
    查找章节内的表格，将人员数据填入。
    """
    personnel_list = []
    for bp in (bid.bid_personnel or []):
        if bp.personnel:
            personnel_list.append({
                'name': bp.personnel.name,
                'role': bp.role or '',
                'title': bp.personnel.title or '',
                'position': bp.personnel.position or '',
                'phone': bp.personnel.phone or '',
                'skills': bp.personnel.skills or '',
            })

    if not personnel_list:
        return

    # 查找该章节内的表格
    body = doc.element.body
    para_count = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_count += 1
        elif tag == 'tbl':
            if content_start <= para_count <= content_end:
                from docx.table import Table
                tbl = Table(child, doc)
                _fill_personnel_table(tbl, personnel_list, project_data)
                return  # 只填充第一个匹配的表格


def _fill_personnel_table(table, personnel_list, project_data):
    """
    向人员表格中填充数据。
    策略：识别表头行，然后向已有的空行填充或复制行模板添加数据。
    """
    if len(table.rows) < 2:
        return

    # 查找数据行的起始位置（跳过表头）
    header_row_idx = 0
    for i, row in enumerate(table.rows):
        cells_text = [c.text.strip() for c in row.cells]
        if any(k in ''.join(cells_text) for k in ['姓名', '序号', '人员']):
            header_row_idx = i
            break

    data_start = header_row_idx + 1
    template_row = table.rows[data_start] if data_start < len(table.rows) else None

    if not template_row:
        return

    # 填充已有行或添加新行
    for idx, person in enumerate(personnel_list):
        row_idx = data_start + idx
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
        else:
            row = _clone_table_row(table, template_row)

        # 尝试按列填充
        cells = row.cells
        col_count = len(cells)
        fill_values = [
            str(idx + 1),           # 序号
            person['name'],         # 姓名
            person['role'],         # 角色/职务
            person['title'],        # 职称
            person['phone'],        # 电话
        ]
        for ci in range(min(col_count, len(fill_values))):
            _set_cell_text_preserve_format(cells[ci], fill_values[ci])


def _fill_performance_section(doc, sec_info, bid, project_data, content_start, content_end):
    """
    填充业绩证明章节。
    查找章节内的表格，将业绩数据填入。
    """
    perf_list = []
    for bp in (bid.bid_performances or []):
        if bp.performance:
            p = bp.performance
            perf_list.append({
                'project_name': p.project_name,
                'client_name': p.client_name or '',
                'amount': f'{p.contract_amount}万元' if p.contract_amount else '',
                'period': '',
                'service_types': p.service_types or '',
            })
            if p.service_start and p.service_end:
                perf_list[-1]['period'] = f'{p.service_start.strftime("%Y.%m")}-{p.service_end.strftime("%Y.%m")}'

    if not perf_list:
        return

    # 查找章节内的表格
    body = doc.element.body
    para_count = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_count += 1
        elif tag == 'tbl':
            if content_start <= para_count <= content_end:
                from docx.table import Table
                tbl = Table(child, doc)
                _fill_performance_table(tbl, perf_list, project_data)
                return


def _fill_performance_table(table, perf_list, project_data):
    """向业绩表格中填充数据"""
    if len(table.rows) < 2:
        return

    header_row_idx = 0
    for i, row in enumerate(table.rows):
        cells_text = [c.text.strip() for c in row.cells]
        if any(k in ''.join(cells_text) for k in ['项目名称', '序号', '业绩']):
            header_row_idx = i
            break

    data_start = header_row_idx + 1
    template_row = table.rows[data_start] if data_start < len(table.rows) else None

    if not template_row:
        return

    for idx, perf in enumerate(perf_list):
        row_idx = data_start + idx
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
        else:
            row = _clone_table_row(table, template_row)

        cells = row.cells
        col_count = len(cells)
        fill_values = [
            str(idx + 1),
            perf['project_name'],
            perf['client_name'],
            perf['amount'],
            perf['period'],
        ]
        for ci in range(min(col_count, len(fill_values))):
            _set_cell_text_preserve_format(cells[ci], fill_values[ci])


def _fill_text_section(doc, sec_info, bid, project_data, content_start, content_end):
    """
    填充文本类章节（如技术方案）。
    如果有自定义内容，在章节内容区域的末尾插入。
    """
    custom_content = sec_info.get('custom_content', '')
    if not custom_content:
        return

    _insert_text_at_position(doc, custom_content, content_start, project_data)


def _fill_custom_content(doc, sec_info, project_data, content_start, content_end):
    """填充自定义文本内容"""
    custom_content = sec_info.get('custom_content', '')
    if not custom_content:
        return

    _insert_text_at_position(doc, custom_content, content_start, project_data)


def _insert_text_at_position(doc, text, after_para_index, project_data):
    """
    在指定段落位置之后插入文本内容。
    复制指定位置段落的格式。
    """
    paragraphs = doc.paragraphs
    if after_para_index >= len(paragraphs):
        return

    # 获取参考段落的格式
    ref_para = paragraphs[after_para_index]

    lines = text.strip().split('\n')
    # 从后往前插入，以保持顺序
    insert_after = ref_para._element

    for line in lines:
        line = line.strip()
        if not line:
            continue
        new_para = _clone_paragraph_format(ref_para, line)
        # 替换新段落中的占位符
        for run in new_para.runs:
            new_text = BRACKET_PLACEHOLDER.sub(
                lambda m: _resolve_placeholder(m.group(1), project_data), run.text
            )
            run.text = new_text
        # 在参考段落之后插入
        insert_after.addnext(new_para._element)
        insert_after = new_para._element


def _clone_paragraph_format(ref_para, new_text):
    """
    创建一个新段落，复制参考段落的格式，使用新文本。
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    # 创建新的段落 XML 元素
    new_p = copy.deepcopy(ref_para._element)

    # 清除原有文本 runs，但保留段落属性 (pPr)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    # 添加新的 run，使用参考段落第一个 run 的格式
    new_r = OxmlElement('w:r')
    if ref_para.runs:
        ref_rPr = ref_para.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))

    new_t = OxmlElement('w:t')
    new_t.text = new_text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)

    return Paragraph(new_p, ref_para._element.getparent())


def _clone_table_row(table, template_row):
    """
    复制表格行（保留格式），添加到表格末尾。
    """
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    from docx.table import _Row
    return _Row(new_tr, table)


def _set_cell_text_preserve_format(cell, text):
    """
    设置单元格文本，保留原有格式。
    如果单元格有内容则替换第一个段落文本，否则设置新文本。
    """
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            # 保留第一个 run 的格式，设置新文本
            para.runs[0].text = text
            for run in para.runs[1:]:
                run.text = ''
        else:
            # 段落没有 run，添加一个
            run = para.add_run(text)
    else:
        cell.text = text


def generate_filled_document(format_file_path, bid, sections_db, app_config=None):
    """
    主入口函数：读取格式模板，结合数据库中的章节信息，生成填充后的文档。

    Args:
        format_file_path: 格式模板文件路径
        bid: BidProject 对象（含关联的 bid_personnel, bid_performances, sections）
        sections_db: BidSection 列表（数据库中的章节记录）
        app_config: Flask app.config

    Returns:
        (Document, str): (填充后的文档对象, 输出文件路径)
    """
    # 从数据库章节记录构建填充信息
    sections_info = []
    for sec in sections_db:
        info = {
            'section_type': sec.section_type or '',
            'section_name': sec.section_name,
            'para_index': sec.format_para_index or 0,
            'content_start': (sec.format_para_index or 0) + 1,
            'content_end': 0,  # 需要计算
            'custom_content': sec.custom_content or '',
            'attachment_id': sec.attachment_id,
        }
        sections_info.append(info)

    # 计算每个章节的 content_end
    sections_info.sort(key=lambda s: s['para_index'])
    for i, sec in enumerate(sections_info):
        if i + 1 < len(sections_info):
            sec['content_end'] = sections_info[i + 1]['para_index']
        else:
            # 最后一个章节延伸到文档末尾
            doc_temp = Document(format_file_path)
            sec['content_end'] = len(doc_temp.paragraphs)

    # 填充模板
    filled_doc = fill_template(format_file_path, sections_info, bid, app_config)

    # 保存到输出目录
    output_dir = os.path.join(os.path.dirname(format_file_path), '..', 'output')
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    output_filename = f'bid_{bid.id}_{timestamp}.docx'
    output_path = os.path.join(output_dir, output_filename)
    filled_doc.save(output_path)

    return filled_doc, output_path
