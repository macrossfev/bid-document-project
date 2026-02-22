#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件解析引擎
从 docx/pdf 招标文件中提取投标文件组成要求，生成结构化章节列表。
"""

import re
import os
from docx import Document
import pdfplumber


# ============================================================
# 章节类型关键词映射字典
# ============================================================

SECTION_KEYWORD_MAP = {
    '投标函': {
        'keywords': ['投标函', '投标书'],
        'category': '投标函模板',
    },
    '报价表': {
        'keywords': ['报价', '分项报价', '投标报价', '开标一览表', '投标总价', '工程量清单'],
        'category': '报价表',
    },
    '法定代表人证明': {
        'keywords': ['法定代表人', '法人身份', '法人证明', '法人授权'],
        'category': '法定代表人证明',
    },
    '授权委托书': {
        'keywords': ['授权委托', '委托书', '总公司授权', '授权文件'],
        'category': '授权委托书',
    },
    '营业执照': {
        'keywords': ['营业执照', '企业法人', '工商注册', '统一社会信用'],
        'category': '营业执照',
    },
    '投标保证金': {
        'keywords': ['保证金', '投标担保', '保函'],
        'category': '其他',
    },
    '基本情况表': {
        'keywords': ['基本情况', '投标人情况', '企业概况', '公司简介'],
        'category': '其他',
    },
    'CMA证书': {
        'keywords': ['CMA', 'CNAS', '资质认定', '检验检测资质', '认可证书', '计量认证'],
        'category': 'CMA证书',
    },
    '资质证书': {
        'keywords': ['资质证书', '资质等级', '行业资质', '专业资质'],
        'category': 'CMA证书',
    },
    '信誉承诺书': {
        'keywords': ['信誉', '承诺书', '无违法', '无行贿', '廉洁', '诚信'],
        'category': '信誉承诺书',
    },
    '人员资料': {
        'keywords': ['人员', '项目负责人', '技术负责人', '职称证', '社保', '资格证书',
                     '采样人员', '检测人员', '持证上岗', '人员配置'],
        'category': '人员资料',
    },
    '技术方案': {
        'keywords': ['技术方案', '实施方案', '服务方案', '检测方案', '技术路线',
                     '工作方案', '质量保证方案', '技术措施'],
        'category': '技术方案',
    },
    '技术偏差表': {
        'keywords': ['技术偏差', '偏差表', '技术要求偏差', '响应偏差'],
        'category': '技术方案',
    },
    '业绩证明': {
        'keywords': ['业绩', '类似项目', '合同', '中标通知', '验收报告', '项目经验', '同类项目'],
        'category': '业绩证明',
    },
    '售后服务': {
        'keywords': ['售后', '服务承诺', '质量保证', '公司优势', '服务保障', '应急预案'],
        'category': '其他',
    },
    '财务报表': {
        'keywords': ['财务', '审计报告', '资产负债', '财务报表', '纳税'],
        'category': '其他',
    },
    '安全生产': {
        'keywords': ['安全生产', '安全管理', '安全制度', '安全许可'],
        'category': '其他',
    },
    '质量管理': {
        'keywords': ['质量管理', 'ISO', '体系认证', '管理体系', '质量体系'],
        'category': '其他',
    },
}

# 用于定位"投标文件组成"段落的关键词模式（按优先级排列）
COMPOSITION_PATTERNS = [
    r'投标文件.*(?:组成|构成|包[含括]|内容)',
    r'投标文件.*(?:应当|应|须|需|必须).*(?:包[含括]|由.*组成)',
    r'(?:投标|响应).*文件.*(?:目录|清单)',
    r'(?:投标|响应).*(?:须知|要求).*(?:组成|构成)',
]

# 用于定位"投标文件格式"章节（第六章格式）的模式
FORMAT_CHAPTER_PATTERNS = [
    r'投标文件格式',
    r'投标文件.*格式',
    r'响应文件格式',
]

# 用于识别编号列表项的正则
NUMBERED_ITEM_PATTERNS = [
    r'^[\s]*[（(]\s*[一二三四五六七八九十]+\s*[）)]',    # （一）（二）
    r'^[\s]*[（(]?\s*(\d+)\s*[）).]',          # (1) 1. 1)
    r'^[\s]*([一二三四五六七八九十]+)\s*[、.]',    # 一、 二、
    r'^[\s]*(\d+\.\d+)',                        # 1.1
    r'^[\s]*[①②③④⑤⑥⑦⑧⑨⑩]',                   # ①②③
    r'^[\s]*[\-\•\·\▪]',                        # - • ·
    r'^[\s]*[a-zA-Z]\s*[）).]',                 # a) b.
]

# 大标题模式: 一、xxx部分  二、xxx部分
PART_HEADING_PATTERN = r'^[\s]*[一二三四五六七八九十]+、.*(?:部分|篇)'


def extract_text_from_docx(filepath):
    """从 docx 文件提取段落文本列表"""
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in [p for p in paragraphs[-10:]]:  # 去重
                    paragraphs.append(text)

    return paragraphs


def extract_text_from_pdf(filepath):
    """从 pdf 文件提取段落文本列表"""
    paragraphs = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        paragraphs.append(line)

            # 也提取表格
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    for cell in row:
                        if cell and cell.strip():
                            paragraphs.append(cell.strip())

    return paragraphs


def _is_toc_line(text):
    """判断是否为目录行（如 '第六章  投标文件格式\t59'）"""
    return bool(re.search(r'\t\d+$', text)) or bool(re.match(r'^.*\s{3,}\d+$', text))


def _is_numbered_item(text):
    """检测是否为编号列表项"""
    for pat in NUMBERED_ITEM_PATTERNS:
        if re.match(pat, text):
            return True
    return False


def _is_part_heading(text):
    """检测是否为大标题（一、xxx部分）"""
    return bool(re.match(PART_HEADING_PATTERN, text))


def _clean_item_prefix(text):
    """清理编号前缀"""
    cleaned = re.sub(r'^[\s]*[（(]\s*[一二三四五六七八九十\d]+\s*[）)]\s*', '', text)
    cleaned = re.sub(r'^[\s]*[一二三四五六七八九十]+\s*[、.\s]+', '', cleaned)
    cleaned = re.sub(r'^[\s]*\d+\s*[.)、]\s*', '', cleaned)
    cleaned = re.sub(r'^[\s]*[①②③④⑤⑥⑦⑧⑨⑩]\s*', '', cleaned)
    cleaned = re.sub(r'^[\s]*[\-\•\·\▪]\s*', '', cleaned)
    cleaned = re.sub(r'^[\s]*[a-zA-Z]\s*[）).]\s*', '', cleaned)
    return cleaned.strip()


def find_composition_section(paragraphs):
    """
    定位投标文件组成/格式相关段落。
    策略：
    1. 优先找"第X章 投标文件格式"的实际章节（非目录行），提取其下的大标题和子标题
    2. 其次找"投标文件组成"/"投标文件应包括"等段落
    返回 (start_index, end_index, strategy) strategy='format_chapter' 或 'composition'
    """
    # 策略1：找"投标文件格式"章节标题（跳过目录行）
    format_start = None
    for i, text in enumerate(paragraphs):
        if _is_toc_line(text):
            continue
        for pattern in FORMAT_CHAPTER_PATTERNS:
            if re.search(pattern, text) and re.match(r'^[\s]*(第[一二三四五六七八九十\d]+[章节])', text):
                format_start = i
                break
        if format_start is not None:
            break

    if format_start is not None:
        # 向下扫描到文档末尾或下一个"第X章"
        end_idx = len(paragraphs)
        for j in range(format_start + 1, len(paragraphs)):
            if re.match(r'^[\s]*第[一二三四五六七八九十\d]+[章节]', paragraphs[j]) and not _is_toc_line(paragraphs[j]):
                end_idx = j
                break
        return format_start, end_idx, 'format_chapter'

    # 策略2：找"投标文件组成"段落（跳过目录行）
    for i, text in enumerate(paragraphs):
        if _is_toc_line(text):
            continue
        for pattern in COMPOSITION_PATTERNS:
            if re.search(pattern, text):
                # 向下扫描收集列表项
                end_idx = i + 1
                consecutive_non_item = 0
                for j in range(i + 1, min(i + 60, len(paragraphs))):
                    t = paragraphs[j]
                    is_major = bool(re.match(r'^[\s]*第[一二三四五六七八九十\d]+[章节部分条]', t))
                    if is_major and j > i + 2:
                        break
                    if _is_numbered_item(t):
                        end_idx = j + 1
                        consecutive_non_item = 0
                    else:
                        consecutive_non_item += 1
                        if len(t) < 50:
                            end_idx = j + 1
                        if consecutive_non_item > 5:
                            break
                return i, end_idx, 'composition'

    return None, None, None


def extract_section_items(paragraphs, start_idx, end_idx, strategy='composition'):
    """
    从候选区域提取章节列表项。

    strategy='format_chapter': 从格式章节中提取大标题（一、二、三）和子标题（（一）（二））
    strategy='composition': 从组成段落中提取编号列表项
    """
    if strategy == 'format_chapter':
        return _extract_from_format_chapter(paragraphs, start_idx, end_idx)
    else:
        return _extract_from_composition(paragraphs, start_idx, end_idx)


def _extract_from_format_chapter(paragraphs, start_idx, end_idx):
    """
    从"第六章 投标文件格式"中提取章节结构。
    识别模式：
    - 大标题: 一、投标函部分  二、资格审查部分  三、技术部分  四、商务部分
    - 子标题: （一）投标函  （二）分项报价表
    - 目录行（"目  录" 下的子标题列表，通常是第一次出现）
    """
    items = []
    seen_names = set()
    current_part = ''

    # 先找是否有"目  录"段，它下面会有紧凑的（一）（二）列表
    # 同时收集大标题
    for i in range(start_idx + 1, end_idx):
        text = paragraphs[i]

        # 大标题: 一、投标函部分
        if _is_part_heading(text):
            current_part = _clean_item_prefix(text)
            continue

        # 子标题: （一）投标函 / （二）分项报价表
        if re.match(r'^[\s]*[（(]\s*[一二三四五六七八九十]+\s*[）)]', text):
            cleaned = _clean_item_prefix(text)
            if cleaned and len(cleaned) < 80:
                # 去重：同一名称只取第一次（目录中的）
                key = re.sub(r'[\s（()）]', '', cleaned)
                if key not in seen_names:
                    seen_names.add(key)
                    name = cleaned
                    if current_part:
                        name = cleaned  # 保留原始名称，part信息通过section_type体现
                    items.append(name)

    return items


def _extract_from_composition(paragraphs, start_idx, end_idx):
    """从"投标文件组成"段落提取列表项"""
    items = []
    current_item = None

    for i in range(start_idx + 1, end_idx):
        text = paragraphs[i]

        if _is_numbered_item(text):
            if current_item:
                items.append(current_item)
            cleaned = _clean_item_prefix(text)
            if cleaned:
                current_item = cleaned
        else:
            if current_item and len(text) < 80:
                current_item += ' ' + text

    if current_item:
        items.append(current_item)

    return items


def classify_section(text):
    """
    对一个章节文本进行关键词匹配分类。
    返回 (section_type, category, score)。
    """
    text_lower = text.lower()
    best_type = '其他'
    best_category = '其他'
    best_score = 0

    for sec_type, info in SECTION_KEYWORD_MAP.items():
        score = 0
        for kw in info['keywords']:
            if kw.lower() in text_lower:
                # 关键词越长，权重越高
                score += len(kw)

        if score > best_score:
            best_score = score
            best_type = sec_type
            best_category = info['category']

    # 归一化分数到 0-100
    normalized_score = min(100, best_score * 15) if best_score > 0 else 0

    return best_type, best_category, normalized_score


def parse_tender_file(filepath):
    """
    解析招标文件，提取投标文件组成要求。

    Args:
        filepath: 招标文件路径 (docx 或 pdf)

    Returns:
        dict: {
            'success': bool,
            'sections': [
                {
                    'order': int,
                    'section_name': str,       # 原始提取的名称
                    'section_type': str,       # 匹配到的标准类型
                    'category': str,           # 对应的资料库分类
                    'match_score': float,      # 匹配置信度
                    'original_text': str,      # 原始文本
                }
            ],
            'raw_context': str,    # 命中的原始上下文段落
            'message': str,
        }
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.docx':
        paragraphs = extract_text_from_docx(filepath)
    elif ext == '.pdf':
        paragraphs = extract_text_from_pdf(filepath)
    else:
        return {
            'success': False,
            'sections': [],
            'raw_context': '',
            'message': f'不支持的文件格式: {ext}，请上传 docx 或 pdf 文件',
        }

    if not paragraphs:
        return {
            'success': False,
            'sections': [],
            'raw_context': '',
            'message': '无法从文件中提取文本内容',
        }

    # 定位投标文件组成段落
    start_idx, end_idx, strategy = find_composition_section(paragraphs)

    if start_idx is None:
        return {
            'success': False,
            'sections': [],
            'raw_context': '',
            'message': '未能在招标文件中找到"投标文件组成"或"投标文件格式"相关段落，请手动添加章节',
        }

    # 提取原始上下文（限制长度避免过长）
    context_lines = paragraphs[start_idx:min(start_idx + 30, end_idx)]
    raw_context = '\n'.join(context_lines)

    # 提取章节列表项
    items = extract_section_items(paragraphs, start_idx, end_idx, strategy)

    if not items:
        return {
            'success': False,
            'sections': [],
            'raw_context': raw_context,
            'message': '找到了"投标文件组成"段落，但未能提取出具体章节列表项',
        }

    # 分类每个章节
    sections = []
    for i, item_text in enumerate(items):
        sec_type, category, score = classify_section(item_text)
        sections.append({
            'order': i + 1,
            'section_name': item_text,
            'section_type': sec_type,
            'category': category,
            'match_score': score,
            'original_text': item_text,
        })

    return {
        'success': True,
        'sections': sections,
        'raw_context': raw_context,
        'message': f'成功提取 {len(sections)} 个章节',
    }


def _find_format_split(paragraphs):
    """
    找到招标文件中 要求部分 与 格式部分 的分割点。
    格式部分通常以 "第X章 投标文件格式" 开头。
    返回 (split_index, format_chapter_title) 或 (None, None)。
    """
    format_patterns = [
        r'^[\s]*第[一二三四五六七八九十\d]+[章节]\s*投标文件格式',
        r'^[\s]*第[一二三四五六七八九十\d]+[章节]\s*投标.*格式',
        r'^[\s]*第[一二三四五六七八九十\d]+[章节]\s*响应文件格式',
    ]
    for i, text in enumerate(paragraphs):
        if _is_toc_line(text):
            continue
        for pat in format_patterns:
            if re.match(pat, text):
                return i, text
    return None, None


def _find_requirements_range(paragraphs, format_start):
    """
    确定要求部分的范围。要求部分一般在 投标须知/资格要求/评分标准 等章节中，
    位于格式章节之前。
    返回 (req_start, req_end)。
    """
    # 要求相关章节的关键词
    req_chapter_keywords = [
        '投标须知', '招标公告', '资格', '评分标准', '评审',
        '技术要求', '商务要求', '合同条款', '投标人须知',
    ]
    req_start = None
    # 从头扫描，找第一个要求相关章节
    for i, text in enumerate(paragraphs):
        if _is_toc_line(text):
            continue
        if re.match(r'^[\s]*第[一二三四五六七八九十\d]+[章节]', text):
            for kw in req_chapter_keywords:
                if kw in text:
                    if req_start is None:
                        req_start = i
                    break
    if req_start is None:
        req_start = 0
    req_end = format_start if format_start else len(paragraphs)
    return req_start, req_end


def _match_requirements_to_section(requirements_text, section_name, section_type):
    """
    从要求文本中找出与某个章节相关的要求段落。
    使用章节名和 SECTION_KEYWORD_MAP 中的关键词进行匹配。
    """
    if not requirements_text:
        return ''

    keywords = []
    # 从章节名提取关键词
    name_keywords = re.findall(r'[\u4e00-\u9fff]{2,}', section_name)
    keywords.extend(name_keywords)

    # 从 SECTION_KEYWORD_MAP 获取关键词
    if section_type and section_type in SECTION_KEYWORD_MAP:
        keywords.extend(SECTION_KEYWORD_MAP[section_type]['keywords'])

    if not keywords:
        return ''

    # 将要求文本按段落分割
    req_lines = requirements_text.split('\n')
    matched_lines = []
    for line in req_lines:
        line = line.strip()
        if not line:
            continue
        for kw in keywords:
            if kw in line:
                matched_lines.append(line)
                break

    return '\n'.join(matched_lines[:10])  # 最多返回10行相关要求


def parse_tender_file_dual(filepath):
    """
    双模式解析招标文件：分别提取要求部分和格式部分。

    Args:
        filepath: 招标文件路径 (docx 或 pdf)

    Returns:
        dict: {
            'success': bool,
            'requirements': str,       # 提取的要求文本（前半部分）
            'format_text': str,        # 提取的格式文本（后半部分）
            'sections': [...],         # 从格式部分提取的章节列表
            'section_requirements': {section_name: requirement_text},
            'raw_context': str,
            'message': str,
        }
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.docx':
        paragraphs = extract_text_from_docx(filepath)
    elif ext == '.pdf':
        paragraphs = extract_text_from_pdf(filepath)
    else:
        return {
            'success': False,
            'requirements': '',
            'format_text': '',
            'sections': [],
            'section_requirements': {},
            'raw_context': '',
            'message': f'不支持的文件格式: {ext}，请上传 docx 或 pdf 文件',
        }

    if not paragraphs:
        return {
            'success': False,
            'requirements': '',
            'format_text': '',
            'sections': [],
            'section_requirements': {},
            'raw_context': '',
            'message': '无法从文件中提取文本内容',
        }

    # 尝试找到格式部分的分割点
    split_idx, split_title = _find_format_split(paragraphs)

    if split_idx is not None:
        # 成功找到分割点 — 双模式解析
        req_start, req_end = _find_requirements_range(paragraphs, split_idx)
        requirements_text = '\n'.join(paragraphs[req_start:req_end])

        # 格式部分：从分割点到下一个章节或文档末尾
        format_end = len(paragraphs)
        for j in range(split_idx + 1, len(paragraphs)):
            if re.match(r'^[\s]*第[一二三四五六七八九十\d]+[章节]', paragraphs[j]) \
                    and not _is_toc_line(paragraphs[j]):
                format_end = j
                break
        format_text = '\n'.join(paragraphs[split_idx:format_end])

        # 从格式部分提取章节
        items = extract_section_items(paragraphs, split_idx, format_end, 'format_chapter')
        raw_context = '\n'.join(paragraphs[split_idx:min(split_idx + 30, format_end)])

        if not items:
            # 格式章节找到了但没提取出子项，回退到完整解析
            fallback = parse_tender_file(filepath)
            fallback['requirements'] = requirements_text
            fallback['format_text'] = format_text
            fallback['section_requirements'] = {}
            # 为回退结果中的每个章节匹配要求
            for sec in fallback.get('sections', []):
                req = _match_requirements_to_section(
                    requirements_text, sec['section_name'], sec.get('section_type', ''))
                fallback['section_requirements'][sec['section_name']] = req
            return fallback

        # 分类章节
        sections = []
        section_requirements = {}
        for i, item_text in enumerate(items):
            sec_type, category, score = classify_section(item_text)
            # 为每个章节匹配相关要求
            req = _match_requirements_to_section(requirements_text, item_text, sec_type)
            sections.append({
                'order': i + 1,
                'section_name': item_text,
                'section_type': sec_type,
                'category': category,
                'match_score': score,
                'original_text': item_text,
                'requirement_text': req,
            })
            section_requirements[item_text] = req
            section_requirements[sec_type] = req  # 也按类型索引

        return {
            'success': True,
            'requirements': requirements_text,
            'format_text': format_text,
            'sections': sections,
            'section_requirements': section_requirements,
            'raw_context': raw_context,
            'message': f'双模式解析成功：提取 {len(sections)} 个章节，已分离要求与格式',
        }

    else:
        # 未找到分割点，回退到单模式解析
        fallback = parse_tender_file(filepath)
        fallback['requirements'] = ''
        fallback['format_text'] = ''
        fallback['section_requirements'] = {}
        fallback['message'] = (fallback.get('message', '') +
                               '（未找到格式章节分割点，使用单模式解析）')
        return fallback


# 提供默认章节列表，当解析失败时使用
DEFAULT_SECTIONS = [
    ('投标函', '投标函', '投标函模板'),
    ('分项报价表', '报价表', '报价表'),
    ('法定代表人身份证明及授权委托书', '法定代表人证明', '法定代表人证明'),
    ('营业执照', '营业执照', '营业执照'),
    ('投标保证金', '投标保证金', '其他'),
    ('投标人基本情况表', '基本情况表', '其他'),
    ('资质认定证书', 'CMA证书', 'CMA证书'),
    ('信誉承诺书', '信誉承诺书', '信誉承诺书'),
    ('项目人员资料', '人员资料', '人员资料'),
    ('技术方案', '技术方案', '技术方案'),
    ('类似项目业绩', '业绩证明', '业绩证明'),
    ('售后服务承诺', '售后服务', '其他'),
]


def get_default_sections():
    """返回默认章节列表（当解析失败或用户选择使用默认时）"""
    sections = []
    for i, (name, sec_type, category) in enumerate(DEFAULT_SECTIONS):
        sections.append({
            'order': i + 1,
            'section_name': name,
            'section_type': sec_type,
            'category': category,
            'match_score': 100,
            'original_text': '',
        })
    return sections


# ============================================================
# 独立要求文档解析（用户手动拆分后的要求文件）
# ============================================================

def parse_requirements_file(filepath):
    """
    解析独立的投标要求文档，提取各项要求内容。
    用户手动从招标文件中拆分出的要求部分。

    Args:
        filepath: 要求文档路径 (docx 或 pdf)

    Returns:
        dict: {
            'success': bool,
            'full_text': str,              # 完整的要求文本
            'requirement_blocks': [        # 按章节/条目分块的要求
                {
                    'title': str,          # 章节/条目标题
                    'content': str,        # 具体要求内容
                    'keywords': [str],     # 提取的关键词
                }
            ],
            'message': str,
        }
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.docx':
        paragraphs = extract_text_from_docx(filepath)
    elif ext == '.pdf':
        paragraphs = extract_text_from_pdf(filepath)
    else:
        return {
            'success': False,
            'full_text': '',
            'requirement_blocks': [],
            'message': f'不支持的文件格式: {ext}',
        }

    if not paragraphs:
        return {
            'success': False,
            'full_text': '',
            'requirement_blocks': [],
            'message': '无法从要求文档中提取文本内容',
        }

    full_text = '\n'.join(paragraphs)

    # 按章节/条目分块
    blocks = _split_into_requirement_blocks(paragraphs)

    return {
        'success': True,
        'full_text': full_text,
        'requirement_blocks': blocks,
        'message': f'成功提取 {len(blocks)} 个要求条目',
    }


def _split_into_requirement_blocks(paragraphs):
    """
    将要求文本按章节标题分块。
    识别章节标题（第X章、一、（一）、1.等），收集其下的内容。
    """
    blocks = []
    current_title = ''
    current_lines = []

    # 章节标题模式
    chapter_patterns = [
        r'^[\s]*第[一二三四五六七八九十\d]+[章节条]\s*(.+)',
        r'^[\s]*[一二三四五六七八九十]+\s*[、.．]\s*(.+)',
        r'^[\s]*[（(]\s*[一二三四五六七八九十]+\s*[）)]\s*(.+)',
        r'^[\s]*\d+\s*[、.．]\s*(.+)',
        r'^[\s]*\d+\.\d+\s+(.+)',
    ]

    for text in paragraphs:
        text = text.strip()
        if not text:
            continue

        is_title = False
        for pat in chapter_patterns:
            m = re.match(pat, text)
            if m:
                # 保存上一个块
                if current_title or current_lines:
                    blocks.append(_make_requirement_block(current_title, current_lines))
                current_title = text
                current_lines = []
                is_title = True
                break

        if not is_title:
            current_lines.append(text)

    # 保存最后一个块
    if current_title or current_lines:
        blocks.append(_make_requirement_block(current_title, current_lines))

    return blocks


def _make_requirement_block(title, content_lines):
    """构建要求块，提取关键词"""
    content = '\n'.join(content_lines)

    # 提取关键词：从标题和内容中找中文词组
    keywords = []
    text_for_kw = title + ' ' + content
    # 匹配 SECTION_KEYWORD_MAP 中的关键词
    for sec_type, info in SECTION_KEYWORD_MAP.items():
        for kw in info['keywords']:
            if kw in text_for_kw:
                keywords.append(kw)

    return {
        'title': title,
        'content': content,
        'keywords': list(set(keywords)),
    }


def match_requirements_to_sections(requirements_result, sections):
    """
    将要求文档中的要求条目与格式模板中的章节进行匹配。

    Args:
        requirements_result: parse_requirements_file() 的返回结果
        sections: 格式模板解析出的章节列表（来自 format_parser）

    Returns:
        dict: {section_name: matched_requirement_text}
    """
    if not requirements_result.get('success'):
        return {}

    full_text = requirements_result.get('full_text', '')
    blocks = requirements_result.get('requirement_blocks', [])
    matched = {}

    for sec in sections:
        sec_name = sec.get('section_name', '')
        sec_type = sec.get('section_type', '')

        # 方法1：用已有的 _match_requirements_to_section 函数做全文匹配
        req_text = _match_requirements_to_section(full_text, sec_name, sec_type)

        # 方法2：从分块中找最相关的块
        best_block_score = 0
        best_block_text = ''
        for block in blocks:
            score = _score_block_match(block, sec_name, sec_type)
            if score > best_block_score:
                best_block_score = score
                best_block_text = block['title'] + '\n' + block['content']

        # 合并两种方法的结果
        if best_block_text and best_block_score > 2:
            if req_text:
                # 去重合并
                combined_lines = req_text.split('\n')
                for line in best_block_text.split('\n'):
                    if line.strip() and line.strip() not in [l.strip() for l in combined_lines]:
                        combined_lines.append(line)
                req_text = '\n'.join(combined_lines[:20])
            else:
                req_text = best_block_text

        matched[sec_name] = req_text
        if sec_type:
            matched[sec_type] = req_text

    return matched


def _score_block_match(block, section_name, section_type):
    """计算要求块与章节的匹配度"""
    score = 0
    block_text = block['title'] + ' ' + block['content']

    # 章节名中的关键词在块中出现
    name_keywords = re.findall(r'[\u4e00-\u9fff]{2,}', section_name)
    for kw in name_keywords:
        if kw in block_text:
            score += len(kw)

    # SECTION_KEYWORD_MAP 中的关键词匹配
    if section_type and section_type in SECTION_KEYWORD_MAP:
        for kw in SECTION_KEYWORD_MAP[section_type]['keywords']:
            if kw in block_text:
                score += len(kw)

    return score
